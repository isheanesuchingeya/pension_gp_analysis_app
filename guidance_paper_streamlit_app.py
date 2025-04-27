###########################################################################################################################################################
                                                                      ## Packages##
import time
start_time = time.time()
import pandas as pd 
import numpy as np 
import openpyxl
import os
import streamlit as st
import janitor
import matplotlib.pyplot as plt
from matplotlib.ticker import MaxNLocator
import matplotlib.pyplot as plt
import seaborn as sns
###########################################################################################################################################################
                                                                  ##Display : Front End##
#Section = st.sidebar.radio("Years", ["2019","2020","2021","2022","2023","2024"])

# --- PAGE SETTINGS ---
st.set_page_config(page_title="Pension Fund Dashboard", layout="wide", page_icon="📊")

# --- HIDE STREAMLIT DEFAULT MENU ---
hide_streamlit_style = """
    <style>
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    header {visibility: hidden;}
    </style>
"""
st.markdown(hide_streamlit_style, unsafe_allow_html=True)

# --- SIDEBAR ---
st.sidebar.title("Settings ⚙️")
st.sidebar.subheader("Upload your file")
uploaded_file = st.sidebar.file_uploader("Upload Pension Excel File", type=["xlsx"])
selected_year = st.sidebar.selectbox("Select Year", ["All", 2019, 2020, 2021, 2022, 2023, 2024])

# --- HEADER ---
st.title("Pension Fund Dashboard 📊")
st.markdown("""
Welcome to the Pension Monitoring System! 
Analyze submissions, assessment status, approval trends, and more.
""")
st.divider()
###########################################################################################################################################################
                                                                      ##Data Import & Cleaning##
# File uploader for the user to upload the Excel file
uploaded_file = st.file_uploader("Upload the Pension Excel file", type=["xlsx"])
if uploaded_file is not None:
    # Read the uploaded Excel file
    pendata_combined = pd.read_excel(uploaded_file)

    # Notify successful upload
    st.success("✅ File uploaded successfully!")

    # You can continue with your processing logic here
    # For example, displaying the first few rows of the DataFrame
    st.write("Sample of Imported Data")
    st.write(pendata_combined.head(3))

else:
    st.warning("⚠️ Please upload an Excel file to proceed.")

# Replace all NaN values with '..not populated'
pendata_combined = pendata_combined.fillna('..not populated')

     # Clean the column names using pyjanitor's clean_names() function
pendata_combined = pendata_combined.clean_names()  # Clean the column names (lowercase, replace spaces with underscores)
     # Shorten column names 
pendata_combined.columns = pendata_combined.columns.str.replace("department", "drpt")
pendata_combined.columns = pendata_combined.columns.str.replace("status", "sts")
###########################################################################################################################################################
                                                      ##Analysis##
# List of years
pensionyears_list = [2019, 2020, 2021, 2022, 2023]

# Initialize an empty dictionary to store the counts
pensionyear_counts = {}

# Loop through each year and count how many times it appears in the 'year' column
for pensionyear in pensionyears_list:  
    pensionyear_counts[pensionyear] = len(pendata_combined[pendata_combined['year'] == pensionyear])

# Convert the dictionary to a DataFrame to display as a table
year_count_table = pd.DataFrame(list(pensionyear_counts.items()), columns=['Year', 'Count'])
# Add borders to the table
year_count_table = year_count_table.style.set_table_styles(
    [{'selector': 'table', 'props': [('border', '1px solid black')]},
     {'selector': 'th', 'props': [('border', '1px solid black')]},
     {'selector': 'td', 'props': [('border', '1px solid black')]},
     {'selector': 'tr', 'props': [('border', '1px solid black')]}]
)



# Initialize an empty dictionary to store the counts
pensionyear_counts = {}

# Loop through each year and count how many times it appears in the 'year' column
for pensionyear in pensionyears_list:  
    pensionyear_counts[pensionyear] = len(pendata_combined[pendata_combined['year'] == pensionyear])

# Convert the dictionary to a DataFrame to display as a table
year_count_table = pd.DataFrame(list(pensionyear_counts.items()), columns=['Year', 'Count'])

# Calculate YoY Growth and replace NaN with 0
year_count_table['Growth rate'] = year_count_table['Count'].pct_change() * 100  

# Replace NaN in YoY Growth with 0 before formatting
year_count_table['Growth rate'] = year_count_table['Growth rate'].fillna(0)

# Format YoY Growth as percentages with 2 decimal places
year_count_table['Growth rate'] = year_count_table['Growth rate'].apply(lambda x: f"{x:.2f}%" if x != 0 else "0.00%")

# Add borders to the table
year_count_table_styled = year_count_table.style.set_table_styles(
    [{'selector': 'table', 'props': [('border', '1px solid black')]},
     {'selector': 'th', 'props': [('border', '1px solid black')]},
     {'selector': 'td', 'props': [('border', '1px solid black')]},
     {'selector': 'tr', 'props': [('border', '1px solid black')]}]
)

#############################################################################################
st.subheader("Quick Overview")
with st.container():
    col1, col2 = st.columns(2)

    with col1:
        st.write("### Submissions over Time")
        pensionyears_list = [2019, 2020, 2021, 2022, 2023]
        pensionyear_counts = {year: len(pendata_combined[pendata_combined['year'] == year]) for year in pensionyears_list}
        year_count_table = pd.DataFrame(list(pensionyear_counts.items()), columns=['Year', 'Count'])
        
        fig, ax = plt.subplots(figsize=(8, 5))
        ax.plot(year_count_table['Year'], year_count_table['Count'], marker='o', color='#2E86AB')
        ax.set_title('Expected Number of Submissions')
        ax.set_xlabel('Year')
        ax.set_ylabel('Count')
        ax.grid(True)
        st.pyplot(fig)

    with col2:
      st.write("### Submission Status Distribution")
      fig, ax = plt.subplots(figsize=(6, 6))
  
      # Count and plot pie chart
      counts = pendata_combined['submission_sts_'].value_counts()
      explode = [0.1] * len(counts)  # Slightly explode all slices for clarity
      
      # Create pie chart
      counts.plot.pie(
          autopct='%1.1f%%',
          colors=['#66b3ff', '#99ff99', '#ffcc99', '#ff9999', '#ffccff'],
          ax=ax, 
          startangle=90,
          explode=explode,
          textprops={'fontsize': 12}  # Increase font size
      )
      
      ax.set_ylabel('')  # Remove the default label
      ax.set_title('Submission Status Distribution', fontsize=14)  # Add a title
      ax.legend(counts.index, title='Status', bbox_to_anchor=(1.05, 1), loc='upper left')  # Add legend
      st.pyplot(fig)

st.subheader("Assessment and Approval Status")
with st.container():
    col3, col4 = st.columns(2)

    with col3:
        fig, ax = plt.subplots(figsize=(8, 5))
        sns.countplot(data=pendata_combined, x='assessment_sts', palette='pastel', ax=ax)
        ax.set_title('Assessment Status')
        ax.set_xlabel('Assessment')
        ax.set_ylabel('Count')
        st.pyplot(fig)

    with col4:
        fig, ax = plt.subplots(figsize=(8, 5))
        sns.countplot(data=pendata_combined, x='approval_sts', palette='pastel', ax=ax)
        ax.set_title('Approval Status')
        ax.set_xlabel('Approval')
        ax.set_ylabel('Count')
        st.pyplot(fig)

#################################################################

# Initialize an empty dictionary to store counts for each status
status_counts = {}
# List of status types we're interested in
status_types = ['Done', 'Partial', 'Pending', 'Troubled', 'Dissolution', 'New fund', 'exemption']
# Loop through each year and count the occurrences of each status
for year in pendata_combined['year'].unique():
    year_data = pendata_combined[pendata_combined['year'] == year]  
    # Initialize a dictionary to store counts for this year
    year_status_counts = {
        'Year': year,
        'Submitted': len(year_data[year_data['submission_sts_'] == 'Done']),
        'Partial': len(year_data[year_data['submission_sts_'] == 'Partial']),
        'Pending': len(year_data[year_data['submission_sts_'] == 'Pending']),
        'Dissolution': len(year_data[year_data['submission_sts_'] == 'Dissolution']),
        'New fund': len(year_data[year_data['submission_sts_'] == 'New fund']),
        'Exemption': len(year_data[year_data['submission_sts_'] == 'exemption']),
        'Troubled': len(year_data[year_data['submission_sts_'] == 'Troubled']),
        'Other': len(year_data[~year_data['submission_sts_'].isin(status_types)])
    }
    
    status_counts[year] = year_status_counts



# Convert the status_counts dictionary to a DataFrame
status_table = pd.DataFrame(status_counts).T  # .T to transpose the dictionary to a DataFrame
# Reset index and drop the default index column that is added when creating a DataFrame
status_table = status_table.reset_index(drop=True)

# Replace NaN values with 0
status_table = status_table.fillna(0)

# Convert all columns to integers (whole numbers)
status_table = status_table.astype(int)
# Remove the row where all values are zero
status_table = status_table[(status_table != 0).any(axis=1)]

# Add borders to the table
status_table_styled = status_table.style.set_table_styles(
    [{'selector': 'table', 'props': [('border', '1px solid black')]},
     {'selector': 'th', 'props': [('border', '1px solid black')]},
     {'selector': 'td', 'props': [('border', '1px solid black')]},
     {'selector': 'tr', 'props': [('border', '1px solid black')]}]
)


# Set the style for seaborn (optional but for better visuals)
sns.set(style="whitegrid")

# Create a new DataFrame that is suitable for a bar plot
status_table_long = status_table.melt(id_vars='Year', var_name='Status', value_name='Count')

# Create a bar plot
plt.figure(figsize=(12, 8))
bar_plot = sns.barplot(data=status_table_long, x='Year', y='Count', hue='Status')

# Customize the plot
plt.title('Submission Status Counts per Year', fontsize=16)
plt.xlabel('Year', fontsize=12)
plt.ylabel('Count', fontsize=12)
plt.legend(title='Status', loc='upper right')
plt.xticks(rotation=45)
plt.tight_layout()
st.pyplot(plt.gcf())


# Set the style for seaborn
sns.set(style="whitegrid")

# Filter the status types to only include 'Done', 'Pending', and 'Partial'
filtered_status_types = ['Submitted', 'Pending', 'Partial']

# Ensure the columns exist before proceeding
if all(status in status_table.columns for status in filtered_status_types):
    # Melt the DataFrame to long format, but only keep the filtered statuses
    status_table_long_filtered = status_table.melt(id_vars='Year', 
                                                   value_vars=filtered_status_types, 
                                                   var_name='Status', 
                                                   value_name='Count')

    # Create the bar plot
    plt.figure(figsize=(12, 8))
    ax = sns.barplot(data=status_table_long_filtered, x='Year', y='Count', hue='Status')


# Define custom colors for the status types
    color_palette = ['#1f77b4', '#d62728', '#ff7f0e']  # blue, red, and yellow
    # Customize the plot
    plt.title('Submission Status Counts for Done, Pending, and Partial per Year', fontsize=16)
    plt.xlabel('Year', fontsize=12)
    plt.ylabel('Count', fontsize=12)
    plt.legend(title='Status', loc='upper right')
    plt.xticks(rotation=45)

    # Add count numbers on top of each bar
    for p in ax.patches:
        ax.annotate(f'{p.get_height()}', 
                    (p.get_x() + p.get_width() / 2., p.get_height()), 
                    ha='center', va='center', 
                    fontsize=10, color='black', 
                    xytext=(0, 5), textcoords='offset points')

    # Show the plot
    plt.tight_layout()
    st.pyplot(plt.gcf())
else:
    print("Some of the status columns are missing from the DataFrame.")


# Define the statuses we want to plot
statuses = [ 'Dissolution','Other', 'Exemption','New fund', 'Troubled',]
status_counts_by_year = {}

# Loop through each year and extract counts for each status
for year in status_table['Year'].unique():
    year_data = status_table[status_table['Year'] == year]
    
    # For each status, count how many times it's greater than 0 for that year
    year_counts = [year_data[status].sum() for status in statuses]
    status_counts_by_year[year] = year_counts

# Convert status counts into a DataFrame
import pandas as pd
status_counts_df = pd.DataFrame(status_counts_by_year, index=statuses).T

# Create a figure and axis for the bar chart
plt.figure(figsize=(12, 8))

# Set positions for each bar
bar_width = 0.15
index = np.arange(len(status_counts_df))

# Define colors for the statuses
colors = ['#1f77b4', '#ff1493', '#32cd32', '#ff7f0e', '#d62728']

# Plot bars for each status type (each status will have bars for each year)
for i, status in enumerate(statuses):
    plt.bar(index + i * bar_width, status_counts_df[status], 
            bar_width, label=status, color=colors[i])

# Customize the plot
plt.title('Submission Status Counts: Dissolution Other, New Fund, Exemption, Troubled', fontsize=16)
plt.xlabel('Year', fontsize=12)
plt.ylabel('Count', fontsize=12)
plt.xticks(index + bar_width * 2, status_counts_df.index, rotation=45)  # Position labels at the center of the grouped bars
plt.legend(title='Status', loc='upper right')

# Add count numbers on top of each bar
for i, status in enumerate(statuses):
    for j, count in enumerate(status_counts_df[status]):
        plt.text(index[j] + i * bar_width, count + 0.5,  # Offset for the text
                 f'{count}', ha='center', va='bottom', fontsize=10)

# Adjust layout and display the plot
plt.tight_layout()
st.pyplot(plt.gcf())


# Initialize an empty dictionary to store counts for each status
assessmentstatus_counts = {}

# List of status types we're interested in
assessmentstatus_types = ['Done', 'Pending', 'Troubled', 'Dissolution', 'New fund', 'exemption']

# Loop through each year and count the occurrences of each status
for year in pendata_combined['year'].unique():
    year_data = pendata_combined[pendata_combined['year'] == year]
    
    # Initialize a dictionary to store counts for this year
    year_assessmentstatus_counts = {
        'Year': year,
        'Assessed': len(year_data[year_data['assessment_sts'] == 'Done']),
        'Non assessed': len(year_data[year_data['assessment_sts'] == 'Pending']),
        'Partial': len(year_data[year_data['assessment_sts'] == 'Partial']),
        'Non assessed': len(year_data[year_data['assessment_sts'] == 'Pending']),
        'Dissolution': len(year_data[year_data['assessment_sts'] == 'Dissolution']),
        'New fund': len(year_data[year_data['assessment_sts'] == 'New fund']),
        'Exemption': len(year_data[year_data['assessment_sts'] == 'exemption']),
        'Troubled': len(year_data[year_data['assessment_sts'] == 'Troubled']),
        'Other': len(year_data[~year_data['assessment_sts'].isin(assessmentstatus_types)])
    }
    
    assessmentstatus_counts[year] = year_assessmentstatus_counts

# Convert the status_counts dictionary to a DataFrame
assessmentstatus_table = pd.DataFrame(assessmentstatus_counts).T  # .T to transpose the dictionary to a DataFrame
# Reset index and drop the default index column that is added when creating a DataFrame
assessmentstatus_table = assessmentstatus_table.reset_index(drop=True)

# Replace NaN values with 0
assessmentstatus_table = assessmentstatus_table.fillna(0)

# Convert all columns to integers (whole numbers)
assessmentstatus_table = assessmentstatus_table.astype(int)

# Remove the row where all values are zero
assessmentstatus_table = assessmentstatus_table[(assessmentstatus_table != 0).any(axis=1)]

# Add borders to the table
assessmentstatus_table_styled = assessmentstatus_table.style.set_table_styles(
    [{'selector': 'table', 'props': [('border', '1px solid black')]},
     {'selector': 'th', 'props': [('border', '1px solid black')]},
     {'selector': 'td', 'props': [('border', '1px solid black')]},
     {'selector': 'tr', 'props': [('border', '1px solid black')]}]
)



# Set the style for seaborn (optional but for better visuals)
sns.set(style="whitegrid")

# Create a new DataFrame that is suitable for a bar plot
# Ensure you use the correct variable name ('assessmentstatus_table')
assessmentstatus_table_long = assessmentstatus_table.melt(id_vars='Year', var_name='Status', value_name='Count')

# Create a bar plot
plt.figure(figsize=(12, 8))
bar_plot = sns.barplot(data=assessmentstatus_table_long, x='Year', y='Count', hue='Status')

# Customize the plot
plt.title('Assessment Status Counts per Year', fontsize=16)
plt.xlabel('Year', fontsize=12)
plt.ylabel('Count', fontsize=12)
plt.legend(title='Status', loc='upper right')
plt.xticks(rotation=45)

# Show the plot
plt.tight_layout()
st.pyplot(plt.gcf())




# Set the style for seaborn
sns.set(style="whitegrid")

# Filter the status types to only include 'Assessed' and 'Non Assessed'
filtered_assessmentstatus_types = ['Assessed', 'Non assessed']  # Ensure this matches the column names exactly

# Ensure the columns exist before proceeding
if all(status in assessmentstatus_table.columns for status in filtered_assessmentstatus_types):
    # Melt the DataFrame to long format, but only keep the filtered statuses
    assessmentstatus_table_long_filtered = assessmentstatus_table.melt(id_vars='Year', 
                                                   value_vars=filtered_assessmentstatus_types, 
                                                   var_name='Status', 
                                                   value_name='Count')

    # Create the bar plot
    plt.figure(figsize=(12, 8))
    ax = sns.barplot(data=assessmentstatus_table_long_filtered, x='Year', y='Count', hue='Status', palette=['#1f77b4', '#ff7f0e'])  # Adjusted the palette to two colors

    # Customize the plot
    plt.title('Assessment Status Counts for Assessed, Non Assessed', fontsize=16)
    plt.xlabel('Year', fontsize=12)
    plt.ylabel('Count', fontsize=12)
    plt.legend(title='Status', loc='upper right')
    plt.xticks(rotation=45)

    # Add count numbers on top of each bar
    for p in ax.patches:
        ax.annotate(f'{p.get_height()}', 
                    (p.get_x() + p.get_width() / 2., p.get_height()), 
                    ha='center', va='center', 
                    fontsize=10, color='black', 
                    xytext=(0, 5), textcoords='offset points')

    # Show the plot
    plt.tight_layout()
    st.pyplot(plt.gcf())
else:
    print("Some of the status columns are missing from the DataFrame.")


import pandas as pd

# Initialize an empty dictionary to store counts for each status
approvalstatus_counts = {}

# List of status types we're interested in
approvalstatus_types = ['Yes', 'No', 'pending']

# Loop through each year and count the occurrences of each status
for year in pendata_combined['year'].unique():
    year_data = pendata_combined[pendata_combined['year'] == year]
    
    # Initialize a dictionary to store counts for this year
    year_approvalstatus_counts = {
        'Year': year,
        'Approved': len(year_data[year_data['approval_sts'] == 'Yes']),
        'Non approved': len(year_data[year_data['approval_sts'] == 'No']),  
        'Pending': len(year_data[year_data['approval_sts'] == 'pending']),  
        'Not populated': len(year_data[~year_data['approval_sts'].isin(approvalstatus_types)]),  # Corrected to check if status is not in the list
    }
    
    # Store the counts for this year
    approvalstatus_counts[year] = year_approvalstatus_counts

# Convert the approvalstatus_counts dictionary to a DataFrame
approvalstatus_table = pd.DataFrame(approvalstatus_counts).T  # .T to transpose the dictionary to a DataFrame
# Reset index and drop the default index column that is added when creating a DataFrame
approvalstatus_table = approvalstatus_table.reset_index(drop=True)

# Replace NaN values with 0
approvalstatus_table = approvalstatus_table.fillna(0)

# Convert all columns to integers (whole numbers)
approvalstatus_table = approvalstatus_table.astype(int)
#Remove rows where 'Year' is 0 (i.e., blank rows)
approvalstatus_table = approvalstatus_table[approvalstatus_table['Year'] != 0]
# Remove the row where all values are zero
approvalstatus_table = approvalstatus_table[(approvalstatus_table != 0).any(axis=1)]

# Add borders to the table for better visibility
approvalstatus_table_styled = approvalstatus_table.style.set_table_styles(
    [{'selector': 'table', 'props': [('border', '1px solid black')]},
     {'selector': 'th', 'props': [('border', '1px solid black')]},
     {'selector': 'td', 'props': [('border', '1px solid black')]},
     {'selector': 'tr', 'props': [('border', '1px solid black')]}]
)


# Set the style for seaborn
sns.set(style="whitegrid")

# Filter the status types to include 'Approved', 'Non approved', and 'Not yet Assessed'
filtered_approvalstatus_types = ['Approved', 'Non approved','Pending', 'Not populated']  # Ensure this matches the column names exactly

# Ensure the columns exist before proceeding
if all(status in approvalstatus_table.columns for status in filtered_approvalstatus_types):
    # Melt the DataFrame to long format, but only keep the filtered statuses
    approvalstatus_table_long_filtered = approvalstatus_table.melt(id_vars='Year', 
                                                                  value_vars=filtered_approvalstatus_types, 
                                                                  var_name='Status', 
                                                                  value_name='Count')

    # Create the bar plot
    plt.figure(figsize=(12, 8))
    ax = sns.barplot(data=approvalstatus_table_long_filtered, x='Year', y='Count', hue='Status', 
                     palette=['#1f77b4', '#ff7f0e', '#2ca02c','#d62728',])  # Adjusted the palette to three colors

    # Customize the plot
    plt.title('Approval Status Counts for Approved, Non Approved,Pending and Not populated', fontsize=16)
    plt.xlabel('Year', fontsize=12)
    plt.ylabel('Count', fontsize=12)
    plt.legend(title='Status', loc='upper right')
    plt.xticks(rotation=45)

    # Add count numbers on top of each bar
    for p in ax.patches:
        ax.annotate(f'{p.get_height()}', 
                    (p.get_x() + p.get_width() / 2., p.get_height()), 
                    ha='center', va='center', 
                    fontsize=10, color='black', 
                    xytext=(0, 5), textcoords='offset points')

    # Show the plot
    plt.tight_layout()
    st.pyplot(plt.gcf())
else:
    print("Some of the approval status columns are missing from the DataFrame.")


approvalstatus_table['Growth (%)'] = approvalstatus_table['Approved'].pct_change() * 100
# Replace NaN in YoY Growth with 0 before formatting
approvalstatus_table['Growth (%)'] = approvalstatus_table['Growth (%)'].fillna(0)


approvalstatus_table['Proportion Approved'] = approvalstatus_table['Approved'] / approvalstatus_table['Approved'].sum()
sns.barplot(data=approvalstatus_table, x='Year', y='Proportion Approved', color='blue')


# Loop through each column to handle NaN replacement
for col in pendata_combined.columns:
    # Check if the column can be converted to numeric
    if pd.api.types.is_numeric_dtype(pendata_combined[col]):
        # Fill NaN in numeric columns with 0
        pendata_combined[col].fillna(0, inplace=True)
    else:
        # Fill NaN in non-numeric columns with 'Not Populated'
        pendata_combined[col].fillna('Not Populated', inplace=True)




pendata_combined.loc[pendata_combined.year==2019]



# Filter data for 2019
pendata_2019 = pendata_combined[pendata_combined['year'] == 2019]


# Map the submission statuses to new categories
pendata_2019['submission_status'] = pendata_2019['submission_sts_'].apply(
    lambda x: 'Submitted' if x == 'Done' else ('Partial' if x == 'Partial' else ('Pending' if x == 'Pending' else 'Other'))

)


# Create a table of submission status counts
submission_status_table2019 = pendata_2019['submission_status'].value_counts().reset_index()
submission_status_table2019.columns = ['Submission Status', 'Count']
print(submission_status_table2019)


# Create a table of submission status counts
submission_status_table2019 = pendata_2019['submission_status'].value_counts().reset_index()
submission_status_table2019.columns = ['Submission Status', 'Count']


fig, ax = plt.subplots(figsize=(6, 3))
ax.axis('tight')
ax.axis('off')
table = ax.table(cellText=submission_status_table2019.values, colLabels=submission_status_table2019.columns, loc='center', cellLoc='center')

# Add borders around the table
for (i, j), cell in table.get_celld().items():
    cell.set_edgecolor('black')  # Set the border color
    cell.set_linewidth(1)        # Set the border thickness
st.pyplot(plt.gcf())


import matplotlib.pyplot as plt

# Bar chart with different colors and count labels
colors = ['#66b3ff', '#99ff99', '#ffcc99']  # Define different colors for each bar
fig, ax = plt.subplots(figsize=(8, 5))

# Plot the bar chart with custom colors
bars = ax.bar(submission_status_table2019['Submission Status'], submission_status_table2019['Count'], color=colors)

# Add title and labels
ax.set_title('Submission Status Distribution (2019)', fontsize=14)
ax.set_xlabel('Submission Status', fontsize=12)
ax.set_ylabel('Count', fontsize=12)

# Add count labels on the bars
for bar in bars:
    yval = bar.get_height()
    ax.text(bar.get_x() + bar.get_width()/2, yval + 0.5, str(int(yval)), ha='center', va='bottom', fontsize=12)

# Show the plot
plt.tight_layout()
st.pyplot(plt.gcf())

# Pie chart
plt.figure(figsize=(6, 6))
pendata_2019['submission_status'].value_counts().plot(kind='pie', autopct='%1.1f%%', colors=['#66b3ff', '#99ff99', '#ffcc99'], startangle=90)
plt.title('Submission Status Distribution (2019)', fontsize=14)
plt.ylabel('')  # Remove the ylabel
plt.tight_layout()
st.pyplot(plt.gcf())


# Group by 'administrator' and 'submission_status', then unstack to reshape for the bar chart
administrator_submission2019 = pendata_2019.groupby(['administrator', 'submission_status']).size().unstack().fillna(0)

# Plotting a bar chart (not stacked)
administrator_submission2019.plot(kind='bar', figsize=(10, 6))

# Set the title and labels
plt.title('Submission Status by Administrator', fontsize=14)
plt.xlabel('Administrator', fontsize=12)
plt.ylabel('Count', fontsize=12)

# Rotate x-axis labels for better readability
plt.xticks(rotation=45)

# Show the plot
plt.tight_layout()  # Adjust layout for better spacing
st.pyplot(plt.gcf())

# Group by 'name_of_fund', 'sts' and 'submission_sts_', then unstack to reshape for the bar chart
fund_status_submission2019 = pendata_2019.groupby([ 'sts', 'submission_sts_']).size().unstack().fillna(0)

# Plotting a bar chart (not stacked)
ax = fund_status_submission2019.plot(kind='bar', figsize=(12, 6))

# Set the title and labels
plt.title('Submission Status by Fund and Status', fontsize=14)
plt.xlabel('Fund Status', fontsize=12)
plt.ylabel('Count', fontsize=12)

# Rotate x-axis labels for better readability
plt.xticks(rotation=45)

# Add count numbers on top of the bars
for container in ax.containers:
    ax.bar_label(container, label_type='edge', fontsize=10, color='black')

# Show the plot
plt.tight_layout()  # Adjust layout for better spacing
st.pyplot(plt.gcf())


# Classify the assessment status in 2019
pendata_2019['assessment_status_category'] = pendata_2019['assessment_sts'].apply(
    lambda x: 'Assessed' if x == 'Done' else ('Pending' if x == 'Pending' else 'Other')
)


# Perform counts for assessment status categories in 2019
assessment_status_table_2019 = pendata_2019['assessment_status_category'].value_counts().reset_index()
assessment_status_table_2019.columns = ['Assessment Status', 'Count']

# Display the table with borders
assementstatusstyled_table2019 = assessment_status_table_2019.style.set_table_styles(
    [{'selector': 'table', 'props': [('border', '2px solid black')]}, 
     {'selector': 'th', 'props': [('border', '2px solid black')]}, 
     {'selector': 'td', 'props': [('border', '2px solid black')]}]
)

# Display the styled table
assementstatusstyled_table2019 


# Get the counts for each assessment status category
assessment_status_counts_2019 = pendata_2019['assessment_status_category'].value_counts()

# Plotting the bar chart for assessment status distribution
plt.figure(figsize=(10, 6))
bars = plt.bar(assessment_status_counts_2019.index, assessment_status_counts_2019.values, color=['#ff9999', '#66b3ff', '#99ff99'])

# Adding title and labels
plt.title('Assessment Status Distribution (2019)', fontsize=14)
plt.xlabel('Assessment Status', fontsize=12)
plt.ylabel('Count', fontsize=12)

# Adding the count numbers on top of the bars
for bar in bars:
    plt.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.05, 
             f'{int(bar.get_height())}', ha='center', fontsize=10)

# Show the plot
plt.tight_layout()
st.pyplot(plt.gcf())


# In[65]:


import matplotlib.pyplot as plt

# Get the counts for each assessment status category
assessment_status_counts_2019 = pendata_2019['assessment_status_category'].value_counts()

# Plotting the pie chart for assessment status distribution
plt.figure(figsize=(8, 8))
plt.pie(assessment_status_counts_2019, labels=assessment_status_counts_2019.index, autopct='%1.1f%%', 
        startangle=90, colors=['#ff9999', '#66b3ff', '#99ff99'])

# Adding title
plt.title('Assessment Status Distribution (2019)', fontsize=14)

# Show the plot
plt.tight_layout()
st.pyplot(plt.gcf())


# In[66]:


# Count the occurrences of each report location
report_location_counts2019 = pendata_2019['report_location'].value_counts().reset_index()

# Rename the columns for clarity
report_location_counts2019.columns = ['Report Location', 'Count']

# Display the result
print(report_location_counts2019)


# In[67]:


# Apply border styling to the report location count table
reportlocationstyled_table2019 = report_location_counts2019.style.set_table_styles(
    [{'selector': 'table', 'props': [('border', '2px solid black')]}, 
     {'selector': 'th', 'props': [('border', '2px solid black')]}, 
     {'selector': 'td', 'props': [('border', '2px solid black')]}]
)

# Display the styled table
reportlocationstyled_table2019


# In[68]:


# Create a new column for approval status categories based on the given criteria
def categorize_approval_status(status):
    if status == 'Yes':
        return 'Approved'
    elif status == 'pending':
        return 'Pending'
    elif status == 'No':
        return 'No'
    else:
        return 'Not populated'

# Apply the categorization function to the approval status column
pendata_2019['approval_status_category'] = pendata_2019['approval_sts'].apply(categorize_approval_status)




# In[69]:


# Count the occurrences of each approval status category
approval_status_counts2019 = pendata_2019['approval_status_category'].value_counts().reset_index()
approval_status_counts2019.columns = ['Approval Status', 'Count']

# Display the result
print(approval_status_counts2019)


# In[70]:


# Apply border styling to the approval status count table
styled_approval_status_table2019 = approval_status_counts2019.style.set_table_styles(
    [{'selector': 'table', 'props': [('border', '2px solid black')]}, 
     {'selector': 'th', 'props': [('border', '2px solid black')]}, 
     {'selector': 'td', 'props': [('border', '2px solid black')]}]
)

# Display the styled table
styled_approval_status_table2019


# In[71]:


# Plotting the bar chart for approval status categories
plt.figure(figsize=(7, 6))
bars = plt.bar(approval_status_counts2019['Approval Status'], approval_status_counts2019['Count'], color='#ff9999')

# Adding title and labels
plt.title('Count of Approval Status (2019)', fontsize=14)
plt.xlabel('Approval Status', fontsize=12)
plt.ylabel('Count', fontsize=12)

# Adding the count numbers on top of the bars
for bar in bars:
    plt.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.05, 
             f'{int(bar.get_height())}', ha='center', fontsize=10)

# Rotate x-axis labels for readability
plt.xticks(rotation=45, ha='right')

# Show the plot
plt.tight_layout()
st.pyplot(plt.gcf())


# In[72]:


# Plotting the pie chart for approval status categories
plt.figure(figsize=(8, 8))
plt.pie(approval_status_counts2019['Count'], labels=approval_status_counts2019['Approval Status'], autopct='%1.1f%%', colors=['#ff9999', '#66b3ff', '#99ff99'])
plt.title('Approval Status Distribution (2019)', fontsize=14)
st.pyplot(plt.gcf())



# Loop through each column to handle NaN replacement
for col in pendata_combined.columns:
    # Check if the column can be converted to numeric
    if pd.api.types.is_numeric_dtype(pendata_combined[col]):
        # Fill NaN in numeric columns with 0
        pendata_combined[col].fillna(0, inplace=True)
    else:
        # Fill NaN in non-numeric columns with 'Not Populated'
        pendata_combined[col].fillna('Not Populated', inplace=True)




# Loop through each column to handle NaN replacement
for col in pendata_combined.columns:
    # Check if the column can be converted to numeric
    if pd.api.types.is_numeric_dtype(pendata_combined[col]):
        # Fill NaN in numeric columns with 0
        pendata_combined[col].fillna(0, inplace=True)
    else:
        # Fill NaN in non-numeric columns with 'Not Populated'
        pendata_combined[col].fillna('Not Populated', inplace=True)


# Filter data for 2020
pendata_2020 = pendata_combined[pendata_combined['year'] == 2020]
# Map the submission statuses to new categories
pendata_2020['submission_status'] = pendata_2020['submission_sts_'].apply(
    lambda x: 'Submitted' if x == 'Done' else ('Partial' if x == 'Partial' else ('Pending' if x == 'Pending' else 'Other'))

)


# Create a table of submission status counts
submission_status_table2020 = pendata_2020['submission_status'].value_counts().reset_index()
submission_status_table2020.columns = ['Submission Status', 'Count']

# Display the table with borders
fig, ax = plt.subplots(figsize=(6, 3))
ax.axis('tight')
ax.axis('off')
table = ax.table(cellText=submission_status_table2020.values, colLabels=submission_status_table2020.columns, loc='center', cellLoc='center')

# Add borders around the table
for (i, j), cell in table.get_celld().items():
    cell.set_edgecolor('black')  # Set the border color
    cell.set_linewidth(1)        # Set the border thickness
st.pyplot(plt.gcf())


# Display the table with borders
import matplotlib.pyplot as plt

fig, ax = plt.subplots(figsize=(6, 3))
ax.axis('tight')
ax.axis('off')
table = ax.table(cellText=submission_status_table2020.values, colLabels=submission_status_table2020.columns, loc='center', cellLoc='center')

# Add borders around the table
for (i, j), cell in table.get_celld().items():
    cell.set_edgecolor('black')  # Set the border color
    cell.set_linewidth(1)        # Set the border thickness
st.pyplot(plt.gcf())



# Bar chart with different colors and count labels
colors = ['#66b3ff', '#99ff99', '#ffcc99']  # Define different colors for each bar
fig, ax = plt.subplots(figsize=(8, 5))

# Plot the bar chart with custom colors
bars = ax.bar(submission_status_table2020['Submission Status'], submission_status_table2020['Count'], color=colors)

# Add title and labels
ax.set_title('Submission Status Distribution (2020)', fontsize=14)
ax.set_xlabel('Submission Status', fontsize=12)
ax.set_ylabel('Count', fontsize=12)

# Add count labels on the bars
for bar in bars:
    yval = bar.get_height()
    ax.text(bar.get_x() + bar.get_width()/2, yval + 0.5, str(int(yval)), ha='center', va='bottom', fontsize=12)

# Show the plot
plt.tight_layout()
st.pyplot(plt.gcf())


# Pie chart
plt.figure(figsize=(6, 6))
pendata_2020['submission_status'].value_counts().plot(kind='pie', autopct='%1.1f%%', colors=['#66b3ff', '#99ff99', '#ffcc99'], startangle=90)
plt.title('Submission Status Distribution (2020)', fontsize=14)
plt.ylabel('')  # Remove the ylabel
plt.tight_layout()
st.pyplot(plt.gcf())

# Classify the assessment status in 2020
pendata_2020['assessment_status_category'] = pendata_2020['assessment_sts'].apply(
    lambda x: 'Assessed' if x == 'Done' else ('Pending' if x == 'Pending' else 'Other')
)

# Perform counts for assessment status categories in 2020
assessment_status_table_2020 = pendata_2020['assessment_status_category'].value_counts().reset_index()
assessment_status_table_2020.columns = ['Assessment Status', 'Count']

# Display the table with borders
assementstatusstyled_table2020 = assessment_status_table_2020.style.set_table_styles(
    [{'selector': 'table', 'props': [('border', '2px solid black')]}, 
     {'selector': 'th', 'props': [('border', '2px solid black')]}, 
     {'selector': 'td', 'props': [('border', '2px solid black')]}]
)

# Display the styled table
assementstatusstyled_table2020


# In[84]:


# Get the counts for each assessment status category
assessment_status_counts_2020 = pendata_2020['assessment_status_category'].value_counts()

# Plotting the bar chart for assessment status distribution
plt.figure(figsize=(10, 6))
bars = plt.bar(assessment_status_counts_2020.index, assessment_status_counts_2020.values, color=['#ff9999', '#66b3ff', '#99ff99'])

# Adding title and labels
plt.title('Assessment Status Distribution (2020)', fontsize=14)
plt.xlabel('Assessment Status', fontsize=12)
plt.ylabel('Count', fontsize=12)

# Adding the count numbers on top of the bars
for bar in bars:
    plt.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.05, 
             f'{int(bar.get_height())}', ha='center', fontsize=10)

# Show the plot
plt.tight_layout()
st.pyplot(plt.gcf())


# In[85]:


# Plotting the pie chart for assessment status distribution
plt.figure(figsize=(8, 8))
plt.pie(assessment_status_counts_2020, labels=assessment_status_counts_2020.index, autopct='%1.1f%%', 
        startangle=90, colors=['#ff9999', '#66b3ff', '#99ff99'])

# Adding title
plt.title('Assessment Status Distribution (2020)', fontsize=14)

# Show the plot
plt.tight_layout()
st.pyplot(plt.gcf())


# In[86]:


# Count the occurrences of each report location
report_location_counts2020 = pendata_2020['report_location'].value_counts().reset_index()

# Rename the columns for clarity
report_location_counts2020.columns = ['Report Location', 'Count']

# Display the result
print(report_location_counts2020)

# Apply border styling to the report location count table
reportlocationstyled_table2020 = report_location_counts2020.style.set_table_styles(
    [{'selector': 'table', 'props': [('border', '2px solid black')]}, 
     {'selector': 'th', 'props': [('border', '2px solid black')]}, 
     {'selector': 'td', 'props': [('border', '2px solid black')]}]
)

# Display the styled table
reportlocationstyled_table2020


# In[87]:


# Count the occurrences of each report location
report_location_counts2020 = pendata_2020['report_location'].value_counts().reset_index()

# Rename the columns for clarity
report_location_counts2020.columns = ['Report Location', 'Count']

# Display the result
print(report_location_counts2020)

# Apply border styling to the report location count table
reportlocationstyled_table2020 = report_location_counts2020.style.set_table_styles(
    [{'selector': 'table', 'props': [('border', '2px solid black')]}, 
     {'selector': 'th', 'props': [('border', '2px solid black')]}, 
     {'selector': 'td', 'props': [('border', '2px solid black')]}]
)

# Display the styled table
reportlocationstyled_table2020


# In[88]:


# Create a new column for approval status categories based on the given criteria
def categorize_approval_status(status):
    if status == 'Yes':
        return 'Approved'
    elif status == 'pending':
        return 'Pending'
    elif status == 'No':
        return 'Non-approved'
    else:
        return 'Not populated'

# Apply the categorization function to the approval status column
pendata_2020['approval_status_category'] = pendata_2020['approval_sts'].apply(categorize_approval_status)

# Count the occurrences of each approval status category
approval_status_counts2020 = pendata_2020['approval_status_category'].value_counts().reset_index()
approval_status_counts2020.columns = ['Approval Status', 'Count']

# Display the result
print(approval_status_counts2020)

# Apply border styling to the approval status count table
styled_approval_status_table2020 = approval_status_counts2020.style.set_table_styles(
    [{'selector': 'table', 'props': [('border', '2px solid black')]}, 
     {'selector': 'th', 'props': [('border', '2px solid black')]}, 
     {'selector': 'td', 'props': [('border', '2px solid black')]}]
)

# Display the styled table
styled_approval_status_table2020


# In[89]:


# Plotting the bar chart for approval status categories
plt.figure(figsize=(7, 6))
bars = plt.bar(approval_status_counts2020['Approval Status'], approval_status_counts2020['Count'], color='#ff9999')

# Adding title and labels
plt.title('Count of Approval Status (2020)', fontsize=14)
plt.xlabel('Approval Status', fontsize=12)
plt.ylabel('Count', fontsize=12)

# Adding the count numbers on top of the bars
for bar in bars:
    plt.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.05, 
             f'{int(bar.get_height())}', ha='center', fontsize=10)

# Show the plot
plt.tight_layout()
st.pyplot(plt.gcf())


# In[90]:


# Loop through each column to handle NaN replacement
for col in pendata_combined.columns:
    # Check if the column can be converted to numeric
    if pd.api.types.is_numeric_dtype(pendata_combined[col]):
        # Fill NaN in numeric columns with 0
        pendata_combined[col].fillna(0, inplace=True)
    else:
        # Fill NaN in non-numeric columns with 'Not Populated'
        pendata_combined[col].fillna('Not Populated', inplace=True)


# In[91]:


pendata_combined.loc[pendata_combined.year == 2021]


# In[92]:


# Filter data for 2021
pendata_2021 = pendata_combined[pendata_combined['year'] == 2021]


# In[93]:


# Map the submission statuses to new categories
pendata_2021['submission_status'] = pendata_2021['submission_sts_'].apply(
    lambda x: 'Assessed' if x == 'Done' else ('Pending' if x == 'Pending' else 'Other')
)

# Create a table of submission status counts
submission_status_table2021 = pendata_2021['submission_status'].value_counts().reset_index()
submission_status_table2021.columns = ['Submission Status', 'Count']
print(submission_status_table2021)


# In[94]:


# Display the table with borders
import matplotlib.pyplot as plt

fig, ax = plt.subplots(figsize=(6, 3))
ax.axis('tight')
ax.axis('off')
table = ax.table(cellText=submission_status_table2021.values, colLabels=submission_status_table2021.columns, loc='center', cellLoc='center')

# Add borders around the table
for (i, j), cell in table.get_celld().items():
    cell.set_edgecolor('black')  # Set the border color
    cell.set_linewidth(1)        # Set the border thickness
st.pyplot(plt.gcf())


# In[95]:


# Bar chart with different colors and count labels
colors = ['#66b3ff', '#99ff99', '#ffcc99']  # Define different colors for each bar
fig, ax = plt.subplots(figsize=(8, 5))

# Plot the bar chart with custom colors
bars = ax.bar(submission_status_table2021['Submission Status'], submission_status_table2021['Count'], color=colors)

# Add title and labels
ax.set_title('Submission Status Distribution (2021)', fontsize=14)
ax.set_xlabel('Submission Status', fontsize=12)
ax.set_ylabel('Count', fontsize=12)

# Add count labels on the bars
for bar in bars:
    yval = bar.get_height()
    ax.text(bar.get_x() + bar.get_width()/2, yval + 0.5, str(int(yval)), ha='center', va='bottom', fontsize=12)

# Show the plot
plt.tight_layout()
st.pyplot(plt.gcf())


# In[96]:


# Pie chart
plt.figure(figsize=(6, 6))
pendata_2021['submission_status'].value_counts().plot(kind='pie', autopct='%1.1f%%', colors=['#66b3ff', '#99ff99', '#ffcc99'], startangle=90)
plt.title('Submission Status Distribution (2021)', fontsize=14)
plt.ylabel('')  # Remove the ylabel
plt.tight_layout()
st.pyplot(plt.gcf())


# In[97]:


# Group by 'administrator' and 'submission_status', then unstack to reshape for the bar chart
administrator_submission2021 = pendata_2021.groupby(['administrator', 'submission_status']).size().unstack().fillna(0)

# Plotting a bar chart (not stacked)
administrator_submission2021.plot(kind='bar', figsize=(10, 6))

# Set the title and labels
plt.title('Submission Status by Administrator', fontsize=14)
plt.xlabel('Administrator', fontsize=12)
plt.ylabel('Count', fontsize=12)

# Rotate x-axis labels for better readability
plt.xticks(rotation=45)

# Show the plot
plt.tight_layout()  # Adjust layout for better spacing
st.pyplot(plt.gcf())


# In[98]:


# Group by 'name_of_fund', 'sts' and 'submission_sts_', then unstack to reshape for the bar chart
fund_status_submission2021 = pendata_2021.groupby([ 'sts', 'submission_sts_']).size().unstack().fillna(0)

# Plotting a bar chart (not stacked)
ax = fund_status_submission2021.plot(kind='bar', figsize=(12, 6))

# Set the title and labels
plt.title('Submission Status by Fund and Status', fontsize=14)
plt.xlabel('Fund Status', fontsize=12)
plt.ylabel('Count', fontsize=12)

# Rotate x-axis labels for better readability
plt.xticks(rotation=45)

# Add count numbers on top of the bars
for container in ax.containers:
    ax.bar_label(container, label_type='edge', fontsize=10, color='black')

# Show the plot
plt.tight_layout()  # Adjust layout for better spacing
st.pyplot(plt.gcf())


# In[99]:


# Classify the assessment status in 2021
pendata_2021['assessment_status_category'] = pendata_2021['assessment_sts'].apply(
    lambda x: 'Assessed' if x == 'Done' else ('Non-Assessed' if x == 'Pending' else 'Other')
)

# Perform counts for assessment status categories in 2021
assessment_status_table_2021 = pendata_2021['assessment_status_category'].value_counts().reset_index()
assessment_status_table_2021.columns = ['Assessment Status', 'Count']

# Display the table with borders
assementstatusstyled_table2021 = assessment_status_table_2021.style.set_table_styles(
    [{'selector': 'table', 'props': [('border', '2px solid black')]}, 
     {'selector': 'th', 'props': [('border', '2px solid black')]}, 
     {'selector': 'td', 'props': [('border', '2px solid black')]}]
)

# Display the styled table
assementstatusstyled_table2021


# In[100]:


# Get the counts for each assessment status category
assessment_status_counts_2021 = pendata_2021['assessment_status_category'].value_counts()

# Plotting the bar chart for assessment status distribution
plt.figure(figsize=(10, 6))
bars = plt.bar(assessment_status_counts_2021.index, assessment_status_counts_2021.values, color=['#ff9999', '#66b3ff', '#99ff99'])

# Adding title and labels
plt.title('Assessment Status Distribution (2021)', fontsize=14)
plt.xlabel('Assessment Status', fontsize=12)
plt.ylabel('Count', fontsize=12)

# Adding the count numbers on top of the bars
for bar in bars:
    plt.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.05, 
             f'{int(bar.get_height())}', ha='center', fontsize=10)

# Show the plot
plt.tight_layout()
st.pyplot(plt.gcf())


# In[101]:


# Plotting the pie chart for assessment status distribution
plt.figure(figsize=(8, 8))
plt.pie(assessment_status_counts_2021, labels=assessment_status_counts_2021.index, autopct='%1.1f%%', 
        startangle=90, colors=['#ff9999', '#66b3ff', '#99ff99'])

# Adding title
plt.title('Assessment Status Distribution (2021)', fontsize=14)

# Show the plot
plt.tight_layout()
st.pyplot(plt.gcf())


# In[102]:


# Count the occurrences of each report location
report_location_counts2021 = pendata_2021['report_location'].value_counts().reset_index()

# Rename the columns for clarity
report_location_counts2021.columns = ['Report Location', 'Count']

# Display the result
print(report_location_counts2021)

# Apply border styling to the report location count table
reportlocationstyled_table2021 = report_location_counts2021.style.set_table_styles(
    [{'selector': 'table', 'props': [('border', '2px solid black')]}, 
     {'selector': 'th', 'props': [('border', '2px solid black')]}, 
     {'selector': 'td', 'props': [('border', '2px solid black')]}]
)

# Display the styled table
reportlocationstyled_table2021


# In[103]:


# Create a new column for approval status categories based on the given criteria
def categorize_approval_status(status):
    if status == 'Yes':
        return 'Approved'
    elif status == 'Pending':
        return 'Non-approved'
    else:
        return 'Not yet assessed'

# Apply the categorization function to the approval status column
pendata_2021['approval_status_category'] = pendata_2021['approval_sts'].apply(categorize_approval_status)

# Count the occurrences of each approval status category
approval_status_counts2021 = pendata_2021['approval_status_category'].value_counts().reset_index()
approval_status_counts2021.columns = ['Approval Status', 'Count']

# Display the result
print(approval_status_counts2021)

# Apply border styling to the approval status count table
styled_approval_status_table2021 = approval_status_counts2021.style.set_table_styles(
    [{'selector': 'table', 'props': [('border', '2px solid black')]}, 
     {'selector': 'th', 'props': [('border', '2px solid black')]}, 
     {'selector': 'td', 'props': [('border', '2px solid black')]}]
)

# Display the styled table
styled_approval_status_table2021


# In[104]:


# Plotting the bar chart for approval status categories
plt.figure(figsize=(7, 6))
bars = plt.bar(approval_status_counts2021['Approval Status'], approval_status_counts2021['Count'], color='#ff9999')

# Adding title and labels
plt.title('Count of Approval Status (2021)', fontsize=14)
plt.xlabel('Approval Status', fontsize=12)
plt.ylabel('Count', fontsize=12)

# Adding the count numbers on top of the bars
for bar in bars:
    plt.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.05, 
             f'{int(bar.get_height())}', ha='center', fontsize=10)

# Show the plot
plt.tight_layout()
st.pyplot(plt.gcf())


# In[105]:


#2022 2022 2022 2022 2022 2022 2022 2022 2022 2022 2022 2022 2022 2022 


# In[106]:


# Loop through each column to handle NaN replacement
for col in pendata_combined.columns:
    # Check if the column can be converted to numeric
    if pd.api.types.is_numeric_dtype(pendata_combined[col]):
        # Fill NaN in numeric columns with 0
        pendata_combined[col].fillna(0, inplace=True)
    else:
        # Fill NaN in non-numeric columns with 'Not Populated'
        pendata_combined[col].fillna('Not Populated', inplace=True)


# In[107]:


pendata_combined.loc[pendata_combined.year == 2022]


# In[108]:


pendata_combined.columns


# In[109]:


# Filter data for 2022
pendata_2022 = pendata_combined[pendata_combined['year'] == 2022]


# In[110]:


# Map the submission statuses to new categories
pendata_2022['submission_status'] = pendata_2022['submission_sts_'].apply(
    lambda x: 'Submitted' if x == 'Done' else ('Partial' if x == 'Partial' else 'Pending')
)

# Create a table of submission status counts
submission_status_table2022 = pendata_2022['submission_status'].value_counts().reset_index()
submission_status_table2022.columns = ['Submission Status', 'Count']
print(submission_status_table2022)


# In[111]:


# Bar chart with different colors and count labels
colors = ['#66b3ff', '#99ff99', '#ffcc99']  # Define different colors for each bar
fig, ax = plt.subplots(figsize=(8, 5))

# Plot the bar chart with custom colors
bars = ax.bar(submission_status_table2022['Submission Status'], submission_status_table2022['Count'], color=colors)

# Add title and labels
ax.set_title('Submission Status Distribution (2022)', fontsize=14)
ax.set_xlabel('Submission Status', fontsize=12)
ax.set_ylabel('Count', fontsize=12)

# Add count labels on the bars
for bar in bars:
    yval = bar.get_height()
    ax.text(bar.get_x() + bar.get_width()/2, yval + 0.5, str(int(yval)), ha='center', va='bottom', fontsize=12)

# Show the plot
plt.tight_layout()
st.pyplot(plt.gcf())


# In[112]:


# Pie chart
plt.figure(figsize=(6, 6))
pendata_2022['submission_status'].value_counts().plot(kind='pie', autopct='%1.1f%%', colors=['#66b3ff', '#99ff99', '#ffcc99'], startangle=90)
plt.title('Submission Status Distribution (2022)', fontsize=14)
plt.ylabel('')  # Remove the ylabel
plt.tight_layout()
st.pyplot(plt.gcf())


# In[113]:


# Group by 'administrator' and 'submission_status', then unstack to reshape for the bar chart
administrator_submission2022 = pendata_2022.groupby(['administrator', 'submission_status']).size().unstack().fillna(0)

# Plotting a bar chart (not stacked)
administrator_submission2022.plot(kind='bar', figsize=(10, 6))

# Set the title and labels
plt.title('Submission Status by Administrator', fontsize=14)
plt.xlabel('Administrator', fontsize=12)
plt.ylabel('Count', fontsize=12)

# Rotate x-axis labels for better readability
plt.xticks(rotation=45)

# Show the plot
plt.tight_layout()  # Adjust layout for better spacing
st.pyplot(plt.gcf())


# In[114]:


# Group by 'name_of_fund', 'sts' and 'submission_sts_', then unstack to reshape for the bar chart
fund_status_submission2022 = pendata_2022.groupby([ 'sts', 'submission_sts_']).size().unstack().fillna(0)

# Plotting a bar chart (not stacked)
ax = fund_status_submission2022.plot(kind='bar', figsize=(12, 6))

# Set the title and labels
plt.title('Submission Status by Fund and Status', fontsize=14)
plt.xlabel('Fund Status', fontsize=12)
plt.ylabel('Count', fontsize=12)

# Rotate x-axis labels for better readability
plt.xticks(rotation=45)

# Add count numbers on top of the bars
for container in ax.containers:
    ax.bar_label(container, label_type='edge', fontsize=10, color='black')

# Show the plot
plt.tight_layout()  # Adjust layout for better spacing
st.pyplot(plt.gcf())


# In[115]:


# Classify the assessment status in 2022
pendata_2022['assessment_status_category'] = pendata_2022['assessment_sts'].apply(
    lambda x: 'Assessed' if x == 'Done' else ('Pending' if x == 'Pending' else 'Other')
)

# Perform counts for assessment status categories in 2022
assessment_status_table_2022 = pendata_2022['assessment_status_category'].value_counts().reset_index()
assessment_status_table_2022.columns = ['Assessment Status', 'Count']

# Display the table with borders
assementstatusstyled_table2022 = assessment_status_table_2022.style.set_table_styles(
    [{'selector': 'table', 'props': [('border', '2px solid black')]}, 
     {'selector': 'th', 'props': [('border', '2px solid black')]}, 
     {'selector': 'td', 'props': [('border', '2px solid black')]}]
)

# Display the styled table
assementstatusstyled_table2022


# In[116]:


# Get the counts for each assessment status category
assessment_status_counts_2022 = pendata_2022['assessment_status_category'].value_counts()

# Plotting the bar chart for assessment status distribution
plt.figure(figsize=(10, 6))
bars = plt.bar(assessment_status_counts_2022.index, assessment_status_counts_2022.values, color=['#ff9999', '#66b3ff', '#99ff99'])

# Adding title and labels
plt.title('Assessment Status Distribution (2022)', fontsize=14)
plt.xlabel('Assessment Status', fontsize=12)
plt.ylabel('Count', fontsize=12)

# Adding the count numbers on top of the bars
for bar in bars:
    plt.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.05, 
             f'{int(bar.get_height())}', ha='center', fontsize=10)

# Show the plot
plt.tight_layout()
st.pyplot(plt.gcf())


# In[117]:


# Plotting the pie chart for assessment status distribution
plt.figure(figsize=(8, 8))
plt.pie(assessment_status_counts_2022, labels=assessment_status_counts_2022.index, autopct='%1.1f%%', 
        startangle=90, colors=['#ff9999', '#66b3ff', '#99ff99'])

# Adding title
plt.title('Assessment Status Distribution (2022)', fontsize=14)

# Show the plot
plt.tight_layout()
st.pyplot(plt.gcf())


# In[118]:


# Count the occurrences of each report location
report_location_counts2022 = pendata_2022['report_location'].value_counts().reset_index()

# Rename the columns for clarity
report_location_counts2022.columns = ['Report Location', 'Count']

# Display the result
print(report_location_counts2022)

# Apply border styling to the report location count table
reportlocationstyled_table2022 = report_location_counts2022.style.set_table_styles(
    [{'selector': 'table', 'props': [('border', '2px solid black')]}, 
     {'selector': 'th', 'props': [('border', '2px solid black')]}, 
     {'selector': 'td', 'props': [('border', '2px solid black')]}]
)

# Display the styled table
reportlocationstyled_table2022


# In[119]:


# Create a new column for approval status categories based on the given criteria
def categorize_approval_status(status):
    if status == 'Yes':
        return 'Approved'
    elif status == 'pending':
        return 'pending'
    elif status == 'No':
        return 'non approval'
    else:
        return 'Not populated'

# Apply the categorization function to the approval status column
pendata_2022['approval_status_category'] = pendata_2022['approval_sts'].apply(categorize_approval_status)

# Count the occurrences of each approval status category
approval_status_counts2022 = pendata_2022['approval_status_category'].value_counts().reset_index()
approval_status_counts2022.columns = ['Approval Status', 'Count']

# Display the result
print(approval_status_counts2022)

# Apply border styling to the approval status count table
styled_approval_status_table2022 = approval_status_counts2022.style.set_table_styles(
    [{'selector': 'table', 'props': [('border', '2px solid black')]}, 
     {'selector': 'th', 'props': [('border', '2px solid black')]}, 
     {'selector': 'td', 'props': [('border', '2px solid black')]}]
)

# Display the styled table
styled_approval_status_table2022


# In[120]:


# Plotting the bar chart for approval status categories
plt.figure(figsize=(7, 6))
bars = plt.bar(approval_status_counts2022['Approval Status'], approval_status_counts2022['Count'], color='#ff9999')

# Adding title and labels
plt.title('Count of Approval Status (2022)', fontsize=14)
plt.xlabel('Approval Status', fontsize=12)
plt.ylabel('Count', fontsize=12)

# Adding the count numbers on top of the bars
for bar in bars:
    plt.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.05, 
             f'{int(bar.get_height())}', ha='center', fontsize=10)

# Show the plot
plt.tight_layout()
st.pyplot(plt.gcf())


# In[121]:


pendata_combined.loc[pendata_combined.year == 2023]


# In[122]:


pendata_combined.columns


# In[123]:


# Filter data for 2023
pendata_2023 = pendata_combined[pendata_combined['year'] == 2023]


# In[124]:


# Map the submission statuses to new categories
pendata_2023['submission_status'] = pendata_2023['submission_sts_'].apply(
    lambda x: 'Submitted' if x == 'Done' else ('Partial' if x == 'Partial' else 'Pending')
)

# Create a table of submission status counts
submission_status_table2023 = pendata_2023['submission_status'].value_counts().reset_index()
submission_status_table2023.columns = ['Submission Status', 'Count']
print(submission_status_table2023)


# Map the submission statuses to new categories
pendata_2023['submission_status'] = pendata_2023['submission_sts_'].apply(
    lambda x: 'Submitted' if x == 'Done' else ('Partial' if x == 'Partial' else 'Pending')
)

# Create a table of submission status counts
submission_status_table2023 = pendata_2023['submission_status'].value_counts().reset_index()
submission_status_table2023.columns = ['Submission Status', 'Count']
print(submission_status_table2023)


# Display the table with borders
import matplotlib.pyplot as plt

fig, ax = plt.subplots(figsize=(6, 3))
ax.axis('tight')
ax.axis('off')
table = ax.table(cellText=submission_status_table2023.values, colLabels=submission_status_table2023.columns, loc='center', cellLoc='center')

# Add borders around the table
for (i, j), cell in table.get_celld().items():
    cell.set_edgecolor('black')  # Set the border color
    cell.set_linewidth(1)        # Set the border thickness
st.pyplot(plt.gcf())


# In[127]:


# Bar chart with different colors and count labels
colors = ['#66b3ff', '#99ff99', '#ffcc99']  # Define different colors for each bar
fig, ax = plt.subplots(figsize=(8, 5))

# Plot the bar chart with custom colors
bars = ax.bar(submission_status_table2023['Submission Status'], submission_status_table2023['Count'], color=colors)

# Add title and labels
ax.set_title('Submission Status Distribution (2023)', fontsize=14)
ax.set_xlabel('Submission Status', fontsize=12)
ax.set_ylabel('Count', fontsize=12)

# Add count labels on the bars
for bar in bars:
    yval = bar.get_height()
    ax.text(bar.get_x() + bar.get_width()/2, yval + 0.5, str(int(yval)), ha='center', va='bottom', fontsize=12)

# Show the plot
plt.tight_layout()
st.pyplot(plt.gcf())



# Pie chart
plt.figure(figsize=(6, 6))
pendata_2023['submission_status'].value_counts().plot(kind='pie', autopct='%1.1f%%', colors=['#66b3ff', '#99ff99', '#ffcc99'], startangle=90)
plt.title('Submission Status Distribution (2023)', fontsize=14)
plt.ylabel('')  # Remove the ylabel
plt.tight_layout()
st.pyplot(plt.gcf())



# Group by 'administrator' and 'submission_status', then unstack to reshape for the bar chart
administrator_submission2023 = pendata_2023.groupby(['administrator', 'submission_status']).size().unstack().fillna(0)

# Plotting a bar chart (not stacked)
administrator_submission2023.plot(kind='bar', figsize=(10, 6))

# Set the title and labels
plt.title('Submission Status by Administrator', fontsize=14)
plt.xlabel('Administrator', fontsize=12)
plt.ylabel('Count', fontsize=12)

# Rotate x-axis labels for better readability
plt.xticks(rotation=45)

# Show the plot
plt.tight_layout()  # Adjust layout for better spacing
st.pyplot(plt.gcf())


# Group by 'name_of_fund', 'sts' and 'submission_sts_', then unstack to reshape for the bar chart
fund_status_submission2023 = pendata_2023.groupby([ 'sts', 'submission_sts_']).size().unstack().fillna(0)

# Plotting a bar chart (not stacked)
ax = fund_status_submission2023.plot(kind='bar', figsize=(12, 6))

# Set the title and labels
plt.title('Submission Status by Fund and Status', fontsize=14)
plt.xlabel('Fund Status', fontsize=12)
plt.ylabel('Count', fontsize=12)

# Rotate x-axis labels for better readability
plt.xticks(rotation=45)

# Add count numbers on top of the bars
for container in ax.containers:
    ax.bar_label(container, label_type='edge', fontsize=10, color='black')

# Show the plot
plt.tight_layout()  # Adjust layout for better spacing
st.pyplot(plt.gcf())



# Classify the assessment status in 2023
pendata_2023['assessment_status_category'] = pendata_2023['assessment_sts'].apply(
    lambda x: 'Assessed' if x == 'Done' else ('Non-Assessed' if x == 'Pending' else 'Other')
)

# Perform counts for assessment status categories in 2023
assessment_status_table_2023 = pendata_2023['assessment_status_category'].value_counts().reset_index()
assessment_status_table_2023.columns = ['Assessment Status', 'Count']

# Display the table with borders
assementstatusstyled_table2023 = assessment_status_table_2023.style.set_table_styles(
    [{'selector': 'table', 'props': [('border', '2px solid black')]}, 
     {'selector': 'th', 'props': [('border', '2px solid black')]}, 
     {'selector': 'td', 'props': [('border', '2px solid black')]}]
)

# Display the styled table
assementstatusstyled_table2023


# Get the counts for each assessment status category
assessment_status_counts_2023 = pendata_2023['assessment_status_category'].value_counts()

# Plotting the bar chart for assessment status distribution
plt.figure(figsize=(10, 6))
bars = plt.bar(assessment_status_counts_2023.index, assessment_status_counts_2023.values, color=['#ff9999', '#66b3ff', '#99ff99'])

# Adding title and labels
plt.title('Assessment Status Distribution (2023)', fontsize=14)
plt.xlabel('Assessment Status', fontsize=12)
plt.ylabel('Count', fontsize=12)

# Adding the count numbers on top of the bars
for bar in bars:
    plt.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.05, 
             f'{int(bar.get_height())}', ha='center', fontsize=10)

# Show the plot
plt.tight_layout()
st.pyplot(plt.gcf())



# Count the occurrences of each report location
report_location_counts2023 = pendata_2023['report_location'].value_counts().reset_index()

# Rename the columns for clarity
report_location_counts2023.columns = ['Report Location', 'Count']

# Display the result
print(report_location_counts2023)

# Apply border styling to the report location count table
reportlocationstyled_table2023 = report_location_counts2023.style.set_table_styles(
    [{'selector': 'table', 'props': [('border', '2px solid black')]}, 
     {'selector': 'th', 'props': [('border', '2px solid black')]}, 
     {'selector': 'td', 'props': [('border', '2px solid black')]}]
)

# Display the styled table
reportlocationstyled_table2023




# Create a new column for approval status categories based on the given criteria
def categorize_approval_status(status):
    if status == 'Yes':
        return 'Approved'
    elif status == 'pending':
        return 'Pending'
    elif status == 'No':
        return 'Non approval'
    else:
        return 'Not populated'

# Apply the categorization function to the approval status column
pendata_2023['approval_status_category'] = pendata_2023['approval_sts'].apply(categorize_approval_status)

# Count the occurrences of each approval status category
approval_status_counts2023 = pendata_2023['approval_status_category'].value_counts().reset_index()
approval_status_counts2023.columns = ['Approval Status', 'Count']

# Display the result
print(approval_status_counts2023)

# Apply border styling to the approval status count table
styled_approval_status_table2023 = approval_status_counts2023.style.set_table_styles(
    [{'selector': 'table', 'props': [('border', '2px solid black')]}, 
     {'selector': 'th', 'props': [('border', '2px solid black')]}, 
     {'selector': 'td', 'props': [('border', '2px solid black')]}]
)

# Display the styled table
styled_approval_status_table2023




# Plotting the bar chart for approval status categories
plt.figure(figsize=(7, 6))
bars = plt.bar(approval_status_counts2023['Approval Status'], approval_status_counts2023['Count'], color='#ff9999')

# Adding title and labels
plt.title('Count of Approval Status (2023)', fontsize=14)
plt.xlabel('Approval Status', fontsize=12)
plt.ylabel('Count', fontsize=12)

# Adding the count numbers on top of the bars
for bar in bars:
    plt.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.05, 
             f'{int(bar.get_height())}', ha='center', fontsize=10)

# Show the plot
plt.tight_layout()
st.pyplot(plt.gcf())

# Import necessary libraries
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
import io

# Create a new Word Document
doc = Document()
doc.add_heading('Guidance Paper Analysis', level=1)
doc.add_paragraph("This document presents the summary statistics for the guidance paper covering the period from 2019 to 2023. The purpose of this report is to provide a concise overview of key data and trends, tracking progress in submissions, assessments, and approvals. By examining these metrics, the report aims to offer insights into the ongoing developments within these areas, supporting further analysis and decision-making based on the collected information.")
doc.add_heading('Summary Statistics (2019 - 2020)', level=1)

# Loop through each column to handle NaN replacement
for col in pendata_combined.columns:
    if pd.api.types.is_numeric_dtype(pendata_combined[col]):
        pendata_combined[col].fillna(0, inplace=True)
    else:
        pendata_combined[col].fillna('Not Populated', inplace=True)

# Filter data for 2020
pendata_2020 = pendata_combined[pendata_combined['year'] == 2020]

# Map the submission statuses to new categories
pendata_2020['submission_status'] = pendata_2020['submission_sts_'].apply(
    lambda x: 'Submitted' if x == 'Done' else ('Partial' if x == 'Partial' else 'Pending')
)

# Create a table of submission status counts
submission_status_table2020 = pendata_2020['submission_status'].value_counts().reset_index()
submission_status_table2020.columns = ['Submission Status', 'Count']

# Add heading for Submission Status Table
doc.add_heading('Submission Status Distribution Table (2020)', level=2)

# Add paragraph for Submission Status Table
doc.add_paragraph("The table below shows the submission status distribution for the year 2020. This table provides a breakdown of the counts for each submission status category.")

# Add submission status table to the document
table = doc.add_table(rows=submission_status_table2020.shape[0]+1, cols=submission_status_table2020.shape[1])
table.style = 'Table Grid'
for j in range(submission_status_table2020.shape[1]):
    table.cell(0, j).text = submission_status_table2020.columns[j]
for i in range(submission_status_table2020.shape[0]):
    for j in range(submission_status_table2020.shape[1]):
        table.cell(i+1, j).text = str(submission_status_table2020.iat[i, j])

# Add heading for Submission Status Bar Chart
doc.add_heading('Submission Status Distribution Bar Chart (2020)', level=2)

# Add paragraph for Submission Status Bar Chart
doc.add_paragraph("The bar chart below illustrates the submission status distribution for the year 2020. It provides a visual representation of the counts for each submission status category.")

# Create and add bar chart directly to Word (without saving to disk)
fig, ax = plt.subplots(figsize=(4.5, 4.5))  # Adjusted size to fit better in Word
colors = ['#66b3ff', '#99ff99', '#ffcc99']
bars = ax.bar(submission_status_table2020['Submission Status'], submission_status_table2020['Count'], color=colors)
ax.set_title('Submission Status Distribution (2020)', fontsize=14)
ax.set_xlabel('Submission Status', fontsize=12)
ax.set_ylabel('Count', fontsize=12)
for bar in bars:
    yval = bar.get_height()
    ax.text(bar.get_x() + bar.get_width()/2, yval + 0.5, str(int(yval)), ha='center', va='bottom', fontsize=12)
plt.tight_layout()

# Save the plot to a BytesIO object
img_stream = io.BytesIO()
plt.savefig(img_stream, format='png')
img_stream.seek(0)  # Go to the start of the stream

# Add the plot image directly to the Word document
doc.add_picture(img_stream, width=Inches(4.5))
plt.close()

# Add heading for Submission Status Pie Chart
doc.add_heading('Submission Status Distribution Pie Chart (2020)', level=2)

# Add paragraph for Submission Status Pie Chart
doc.add_paragraph("The pie chart below shows the submission status distribution for the year 2020. It visually represents the percentage breakdown of the different submission statuses.")

# Create and add pie chart directly to Word (without saving to disk)
plt.figure(figsize=(4.5, 4.5))  # Adjusted size for better fitting
pendata_2020['submission_status'].value_counts().plot(kind='pie', autopct='%1.1f%%', colors=colors, startangle=90)
plt.title('Submission Status Distribution (2020)', fontsize=14)
plt.ylabel('')
plt.tight_layout()

# Save the pie chart to a BytesIO object
img_stream = io.BytesIO()
plt.savefig(img_stream, format='png')
img_stream.seek(0)  # Go to the start of the stream

# Add the pie chart image directly to the Word document
doc.add_picture(img_stream, width=Inches(4.5))
plt.close()

# Classify the assessment status in 2020
pendata_2020['assessment_status_category'] = pendata_2020['assessment_sts'].apply(
    lambda x: 'Assessed' if x == 'Done' else ('Non-Assessed' if x == 'Pending' else 'Other')
)

# Perform counts for assessment status categories in 2020
assessment_status_table_2020 = pendata_2020['assessment_status_category'].value_counts().reset_index()
assessment_status_table_2020.columns = ['Assessment Status', 'Count']

# Add heading for Assessment Status Table
doc.add_heading('Assessment Status Distribution Table (2020)', level=2)

# Add paragraph for Assessment Status Table
doc.add_paragraph("The table below shows the assessment status distribution for the year 2020. This table provides a breakdown of the counts for each assessment status category.")

# Add assessment status table to the document
table = doc.add_table(rows=assessment_status_table_2020.shape[0]+1, cols=assessment_status_table_2020.shape[1])
table.style = 'Table Grid'
for j in range(assessment_status_table_2020.shape[1]):
    table.cell(0, j).text = assessment_status_table_2020.columns[j]
for i in range(assessment_status_table_2020.shape[0]):
    for j in range(assessment_status_table_2020.shape[1]):
        table.cell(i+1, j).text = str(assessment_status_table_2020.iat[i, j])

# Add heading for Assessment Status Bar Chart
doc.add_heading('Assessment Status Distribution Bar Chart (2020)', level=2)

# Add paragraph for Assessment Status Bar Chart
doc.add_paragraph("The bar chart below illustrates the assessment status distribution for the year 2020. It provides a visual representation of the counts for each assessment status category.")

# Plotting the bar chart for assessment status distribution
plt.figure(figsize=(4.5, 4.5))  # Adjusted size for better fitting
bars = plt.bar(assessment_status_table_2020['Assessment Status'], assessment_status_table_2020['Count'], color=['#ff9999', '#66b3ff', '#99ff99'])
plt.title('Assessment Status Distribution (2020)', fontsize=14)
plt.xlabel('Assessment Status', fontsize=12)
plt.ylabel('Count', fontsize=12)
for bar in bars:
    plt.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.05, f'{int(bar.get_height())}', ha='center', fontsize=10)
plt.tight_layout()

# Save the plot to a BytesIO object
img_stream = io.BytesIO()
plt.savefig(img_stream, format='png')
img_stream.seek(0)  # Go to the start of the stream

# Add the plot image directly to the Word document
doc.add_picture(img_stream, width=Inches(4.5))
plt.close()

# Add heading for Assessment Status Pie Chart
doc.add_heading('Assessment Status Distribution Pie Chart (2020)', level=2)

# Add paragraph for Assessment Status Pie Chart
doc.add_paragraph("The pie chart below shows the assessment status distribution for the year 2020. It visually represents the percentage breakdown of the different assessment statuses.")

# Plotting the pie chart for assessment status distribution
plt.figure(figsize=(4.5, 4.5))  # Adjusted size for better fitting
plt.pie(assessment_status_table_2020['Count'], labels=assessment_status_table_2020['Assessment Status'], 
        autopct='%1.1f%%', startangle=90, colors=['#ff9999', '#66b3ff', '#99ff99'])

# Adding title
plt.title('Assessment Status Distribution (2020)', fontsize=14)

# Save the plot to a BytesIO object
img_stream = io.BytesIO()
plt.savefig(img_stream, format='png')
img_stream.seek(0)  # Go to the start of the stream

# Add the plot image directly to the Word document
doc.add_picture(img_stream, width=Inches(4.5))
plt.close()

# Add heading for Approval Status Table
doc.add_heading('Approval Status Distribution Table (2020)', level=2)

# Add paragraph for Approval Status Table
doc.add_paragraph("The table below shows the approval status distribution for the year 2020. This table provides a breakdown of the counts for each approval status category.")

# Function to categorize approval status
def categorize_approval_status(status):
    if status == 'Yes':
        return 'Approved'
    elif status == 'pending':
        return 'Pending'
    elif status == 'No':
        return 'Non approval'
    else:
        return 'Not yet assessed'

# Apply categorization function
pendata_2020['approval_status_category'] = pendata_2020['approval_sts'].apply(categorize_approval_status)

# Count approval status categories
approval_status_counts2020 = pendata_2020['approval_status_category'].value_counts().reset_index()
approval_status_counts2020.columns = ['Approval Status', 'Count']

# Add approval status table to the document
table = doc.add_table(rows=approval_status_counts2020.shape[0]+1, cols=approval_status_counts2020.shape[1])
table.style = 'Table Grid'
for j in range(approval_status_counts2020.shape[1]):
    table.cell(0, j).text = approval_status_counts2020.columns[j]
for i in range(approval_status_counts2020.shape[0]):
    for j in range(approval_status_counts2020.shape[1]):
        table.cell(i+1, j).text = str(approval_status_counts2020.iat[i, j])

# Add heading for Approval Status Bar Chart
doc.add_heading('Approval Status Distribution Bar Chart (2020)', level=2)

# Add paragraph for Approval Status Bar Chart
doc.add_paragraph("The bar chart below illustrates the approval status distribution for the year 2020. It provides a visual representation of the counts for each approval status category.")

# Plotting the bar chart for approval status distribution
plt.figure(figsize=(4.5, 4.5))  # Adjusted size for better fitting
bars = plt.bar(approval_status_counts2020['Approval Status'], approval_status_counts2020['Count'], color=['#ffcccc', '#cce6ff'])
plt.title('Approval Status Distribution (2020)', fontsize=14)
plt.xlabel('Approval Status', fontsize=12)
plt.ylabel('Count', fontsize=12)
for bar in bars:
    plt.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.05, f'{int(bar.get_height())}', ha='center', fontsize=10)
plt.tight_layout()

# Save the plot to a BytesIO object
img_stream = io.BytesIO()
plt.savefig(img_stream, format='png')
img_stream.seek(0)  # Go to the start of the stream

# Add the plot image directly to the Word document
doc.add_picture(img_stream, width=Inches(4.5))
plt.close()

# Add heading for Approval Status Pie Chart
doc.add_heading('Approval Status Distribution Pie Chart (2020)', level=2)

# Add paragraph for Approval Status Pie Chart
doc.add_paragraph("The pie chart below shows the approval status distribution for the year 2020. It visually represents the percentage breakdown of the different approval statuses.")

# Plotting the pie chart for approval status distribution
plt.figure(figsize=(4.5, 4.5))  # Adjusted size for better fitting
plt.pie(approval_status_counts2020['Count'], labels=approval_status_counts2020['Approval Status'], 
        autopct='%1.1f%%', startangle=90, colors=['#ffcccc', '#cce6ff'])

# Adding title
plt.title('Approval Status Distribution (2020)', fontsize=14)

# Save the plot to a BytesIO object
img_stream = io.BytesIO()
plt.savefig(img_stream, format='png')
img_stream.seek(0)  # Go to the start of the stream

# Add the plot image directly to the Word document
doc.add_picture(img_stream, width=Inches(4.5))
plt.close()

# Save the document: Initital_check

#import io
#doc_io = io.BytesIO()
#doc.save(doc_io)
#doc_io.seek(0)
#st.download_button(
 #   label="📄 Download Word Report",
  #  data=doc_io,
   # file_name="Summary_Statistics_2019_2023_with_detailed_graphs.docx",
    #mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
#)




# Import necessary libraries
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
import io

# Create a new Word Document
doc = Document()
doc.add_heading('Guidance Paper Analysis', level=0)
doc.add_paragraph("This document presents the summary statistics for the guidance paper covering the period from 2019 to 2023. The purpose of this report is to provide a concise overview of key data and trends, tracking progress in submissions, assessments, and approvals. By examining these metrics, the report aims to offer insights into the ongoing developments within these areas, supporting further analysis and decision-making based on the collected information.")

# Function to add data, tables, and charts for each year
def add_yearly_data(year):
    # Filter data for the given year
    pendata_year = pendata_combined[pendata_combined['year'] == year]

    # Mapping the submission statuses
    pendata_year['submission_status'] = pendata_year['submission_sts_'].apply(
        lambda x: 'Submitted' if x == 'Done' else ('Partial' if x == 'Partial' else 'Pending')
    )

    # Create the submission status table
    submission_status_table = pendata_year['submission_status'].value_counts().reset_index()
    submission_status_table.columns = ['Submission Status', 'Count']
    
    # Heading for Submission Status Table
    doc.add_heading(f'Submission Status Distribution Table ({year})', level=2)
    doc.add_paragraph(f"The table below shows the submission status distribution for the year {year}. This table provides a breakdown of the counts for each submission status category.")
    
    # Add submission status table to the document
    table = doc.add_table(rows=submission_status_table.shape[0] + 1, cols=submission_status_table.shape[1])
    table.style = 'Table Grid'
    for j in range(submission_status_table.shape[1]):
        table.cell(0, j).text = submission_status_table.columns[j]
    for i in range(submission_status_table.shape[0]):
        for j in range(submission_status_table.shape[1]):
            table.cell(i + 1, j).text = str(submission_status_table.iat[i, j])
    
    # Heading for Submission Status Bar Chart
    doc.add_heading(f'Submission Status Distribution Bar Chart ({year})', level=2)
    doc.add_paragraph(f"The bar chart below illustrates the submission status distribution for the year {year}. It provides a visual representation of the counts for each submission status category.")
    
    # Create and add bar chart directly to Word (without saving to disk)
    fig, ax = plt.subplots(figsize=(4.5, 4.5))  # Adjusted size to fit better in Word
    colors = ['#66b3ff', '#99ff99', '#ffcc99']
    bars = ax.bar(submission_status_table['Submission Status'], submission_status_table['Count'], color=colors)
    ax.set_title(f'Submission Status Distribution ({year})', fontsize=14)
    ax.set_xlabel('Submission Status', fontsize=12)
    ax.set_ylabel('Count', fontsize=12)
    for bar in bars:
        yval = bar.get_height()
        ax.text(bar.get_x() + bar.get_width() / 2, yval + 0.5, str(int(yval)), ha='center', va='bottom', fontsize=12)
    plt.tight_layout()

    # Save the plot to a BytesIO object
    img_stream = io.BytesIO()
    plt.savefig(img_stream, format='png')
    img_stream.seek(0)  # Go to the start of the stream

    # Add the plot image directly to the Word document
    doc.add_picture(img_stream, width=Inches(4.5))
    plt.close()

    # Add Heading for Submission Status Pie Chart
    doc.add_heading(f'Submission Status Distribution Pie Chart ({year})', level=2)
    doc.add_paragraph(f"The pie chart below shows the submission status distribution for the year {year}. It visually represents the percentage breakdown of the different submission statuses.")

    # Create and add pie chart directly to Word (without saving to disk)
    plt.figure(figsize=(4.5, 4.5))  # Adjusted size for better fitting
    pendata_year['submission_status'].value_counts().plot(kind='pie', autopct='%1.1f%%', colors=colors, startangle=90)
    plt.title(f'Submission Status Distribution ({year})', fontsize=14)
    plt.ylabel('')
    plt.tight_layout()

    # Save the pie chart to a BytesIO object
    img_stream = io.BytesIO()
    plt.savefig(img_stream, format='png')
    img_stream.seek(0)  # Go to the start of the stream

    # Add the pie chart image directly to the Word document
    doc.add_picture(img_stream, width=Inches(4.5))
    plt.close()

    # Classify the assessment status
    pendata_year['assessment_status_category'] = pendata_year['assessment_sts'].apply(
        lambda x: 'Assessed' if x == 'Done' else ('Non-Assessed' if x == 'Pending' else 'Other')
    )

    # Create the assessment status table
    assessment_status_table = pendata_year['assessment_status_category'].value_counts().reset_index()
    assessment_status_table.columns = ['Assessment Status', 'Count']
    
    # Heading for Assessment Status Table
    doc.add_heading(f'Assessment Status Distribution Table ({year})', level=2)
    doc.add_paragraph(f"The table below shows the assessment status distribution for the year {year}. This table provides a breakdown of the counts for each assessment status category.")
    
    # Add assessment status table to the document
    table = doc.add_table(rows=assessment_status_table.shape[0] + 1, cols=assessment_status_table.shape[1])
    table.style = 'Table Grid'
    for j in range(assessment_status_table.shape[1]):
        table.cell(0, j).text = assessment_status_table.columns[j]
    for i in range(assessment_status_table.shape[0]):
        for j in range(assessment_status_table.shape[1]):
            table.cell(i + 1, j).text = str(assessment_status_table.iat[i, j])
    
    # Heading for Assessment Status Bar Chart
    doc.add_heading(f'Assessment Status Distribution Bar Chart ({year})', level=2)
    doc.add_paragraph(f"The bar chart below illustrates the assessment status distribution for the year {year}. It provides a visual representation of the counts for each assessment status category.")
    
    # Create and add bar chart directly to Word (without saving to disk)
    plt.figure(figsize=(4.5, 4.5))  # Adjusted size for better fitting
    bars = plt.bar(assessment_status_table['Assessment Status'], assessment_status_table['Count'], color=['#ff9999', '#66b3ff', '#99ff99'])
    plt.title(f'Assessment Status Distribution ({year})', fontsize=14)
    plt.xlabel('Assessment Status', fontsize=12)
    plt.ylabel('Count', fontsize=12)
    for bar in bars:
        plt.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.05, f'{int(bar.get_height())}', ha='center', fontsize=10)
    plt.tight_layout()

    # Save the plot to a BytesIO object
    img_stream = io.BytesIO()
    plt.savefig(img_stream, format='png')
    img_stream.seek(0)  # Go to the start of the stream

    # Add the plot image directly to the Word document
    doc.add_picture(img_stream, width=Inches(4.5))
    plt.close()

    # Add Heading for Assessment Status Pie Chart
    doc.add_heading(f'Assessment Status Distribution Pie Chart ({year})', level=2)
    doc.add_paragraph(f"The pie chart below shows the assessment status distribution for the year {year}. It visually represents the percentage breakdown of the different assessment statuses.")

    # Create and add pie chart directly to Word (without saving to disk)
    plt.figure(figsize=(4.5, 4.5))  # Adjusted size for better fitting
    plt.pie(assessment_status_table['Count'], labels=assessment_status_table['Assessment Status'], 
            autopct='%1.1f%%', startangle=90, colors=['#ff9999', '#66b3ff', '#99ff99'])
    plt.title(f'Assessment Status Distribution ({year})', fontsize=14)

    # Save the plot to a BytesIO object
    img_stream = io.BytesIO()
    plt.savefig(img_stream, format='png')
    img_stream.seek(0)  # Go to the start of the stream

    # Add the pie chart image directly to the Word document
    doc.add_picture(img_stream, width=Inches(4.5))
    plt.close()

    # Classify the approval status
    def categorize_approval_status(status):
        if status == 'Yes':
            return 'Approved'
        elif status == 'pending':
            return '    Pending'
        else:
            return 'Not populated'

    # Apply categorization function
    pendata_year['approval_status_category'] = pendata_year['approval_sts'].apply(categorize_approval_status)

    # Create the approval status table
    approval_status_table = pendata_year['approval_status_category'].value_counts().reset_index()
    approval_status_table.columns = ['Approval Status', 'Count']
    
    # Heading for Approval Status Table
    doc.add_heading(f'Approval Status Distribution Table ({year})', level=2)
    doc.add_paragraph(f"The table below shows the approval status distribution for the year {year}. This table provides a breakdown of the counts for each approval status category.")
    
    # Add approval status table to the document
    table = doc.add_table(rows=approval_status_table.shape[0] + 1, cols=approval_status_table.shape[1])
    table.style = 'Table Grid'
    for j in range(approval_status_table.shape[1]):
        table.cell(0, j).text = approval_status_table.columns[j]
    for i in range(approval_status_table.shape[0]):
        for j in range(approval_status_table.shape[1]):
            table.cell(i + 1, j).text = str(approval_status_table.iat[i, j])

    # Heading for Approval Status Bar Chart
    doc.add_heading(f'Approval Status Distribution Bar Chart ({year})', level=2)
    doc.add_paragraph(f"The bar chart below illustrates the approval status distribution for the year {year}. It provides a visual representation of the counts for each approval status category.")
    
    # Create and add bar chart directly to Word (without saving to disk)
    plt.figure(figsize=(4.5, 4.5))  # Adjusted size for better fitting
    bars = plt.bar(approval_status_table['Approval Status'], approval_status_table['Count'], color=['#ffcc99', '#66cc99'])
    plt.title(f'Approval Status Distribution ({year})', fontsize=14)
    plt.xlabel('Approval Status', fontsize=12)
    plt.ylabel('Count', fontsize=12)
    for bar in bars:
        plt.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.05, f'{int(bar.get_height())}', ha='center', fontsize=10)
    plt.tight_layout()

    # Save the plot to a BytesIO object
    img_stream = io.BytesIO()
    plt.savefig(img_stream, format='png')
    img_stream.seek(0)  # Go to the start of the stream

    # Add the plot image directly to the Word document
    doc.add_picture(img_stream, width=Inches(4.5))
    plt.close()

    # Add Heading for Approval Status Pie Chart
    doc.add_heading(f'Approval Status Distribution Pie Chart ({year})', level=2)
    doc.add_paragraph(f"The pie chart below shows the approval status distribution for the year {year}. It visually represents the percentage breakdown of the different approval statuses.")

    # Create and add pie chart directly to Word (without saving to disk)
    plt.figure(figsize=(4.5, 4.5))  # Adjusted size for better fitting
    plt.pie(approval_status_table['Count'], labels=approval_status_table['Approval Status'], 
            autopct='%1.1f%%', startangle=90, colors=['#ffcc99', '#66cc99'])
    plt.title(f'Approval Status Distribution ({year})', fontsize=14)

    # Save the plot to a BytesIO object
    img_stream = io.BytesIO()
    plt.savefig(img_stream, format='png')
    img_stream.seek(0)  # Go to the start of the stream

    # Add the pie chart image directly to the Word document
    doc.add_picture(img_stream, width=Inches(4.5))
    plt.close()

# Add data for each year
years = [2019, 2020, 2021, 2022, 2023]
for year in years:
    doc.add_heading(f'Year {year}', level=1)
    add_yearly_data(year)




from io import BytesIO
import matplotlib.pyplot as plt
import seaborn as sns
from docx.shared import Inches

# Add the Year Count Table to Word
doc.add_heading('Year to Year analysis', level=0)

# Add the Year Count Table with YoY Growth
doc.add_heading('Year Count Table with YoY Growth', level=1)

# Add paragraph explaining the table and graph
doc.add_paragraph(
    "The table below shows the number of pension submissions from 2019 to 2023, with a year-over-year (YoY) growth column. "
    "This growth column represents the percentage change in submissions between consecutive years, indicating the "
    "trends in submission frequency over the past five years. The graph below visualizes the expected number of submissions "
    "from 2019 to 2023, clearly showing the changes in the number of submissions year over year. "
    "It allows us to observe whether there has been an increase or decrease in submissions in recent years, which can "
    "offer insights into any shifts in the pension submission process."
)

year_count_table = pd.DataFrame(list(pensionyear_counts.items()), columns=['Year', 'Count'])
year_count_table['Growth rate'] = year_count_table['Count'].pct_change() * 100
year_count_table['Growth rate'] = year_count_table['Growth rate'].fillna(0)
year_count_table['Growth rate'] = year_count_table['Growth rate'].apply(lambda x: f"{x:.2f}%" if x != 0 else "0.00%")

# Add Year Count table to the document using Table Grid style
table = doc.add_table(rows=year_count_table.shape[0] + 1, cols=year_count_table.shape[1])
table.style = 'Table Grid'

# Add header row
for j in range(year_count_table.shape[1]):
    table.cell(0, j).text = year_count_table.columns[j]

# Add table data
for i in range(year_count_table.shape[0]):
    for j in range(year_count_table.shape[1]):
        table.cell(i + 1, j).text = str(year_count_table.iat[i, j])

# Create a plot for the Expected Submissions and insert directly into the Word document
plt.figure(figsize=(4.5, 4.5))
plt.plot(year_count_table['Year'], year_count_table['Count'], marker='o', color='b', linestyle='-', linewidth=2, markersize=6)
plt.title('Expected Number of Submissions', fontsize=14)
plt.xlabel('Year', fontsize=12)
plt.ylabel('Count', fontsize=12)
plt.grid(True)

# Save plot to BytesIO
img_stream = BytesIO()
plt.savefig(img_stream, format='png')
plt.close()
img_stream.seek(0)

# Insert the plot directly into the Word document
doc.add_paragraph(
    "Graph: Expected Number of Submissions. The line graph illustrates the expected number of pension submissions "
    "over the years from 2019 to 2023. It clearly shows the trends in submission frequency, including any significant "
    "peaks or declines. "
    
)
doc.add_picture(img_stream, width=Inches(4.5))

doc.add_paragraph('\n' * 2)

# Add Submission Status Table to Word
doc.add_heading('Submission Status Table by Year', level=1)

# Add paragraph explaining the table and graph
doc.add_paragraph(
    "The table below shows the submission status counts for each year from 2019 to 2023. "
    "The data is broken down into different submission statuses for each year. The graph below illustrates these counts over time, "
    "giving a clear view of how the submission statuses have evolved over the years. By visualizing the data, we can observe patterns "
    "and assess whether certain submission statuses have become more or less prevalent, shedding light on any issues or improvements in the submission process."
)

status_table = pd.DataFrame(status_counts)
status_table = status_table.T.fillna(0).astype(int)

# Add Submission Status table using Table Grid style
table = doc.add_table(rows=status_table.shape[0] + 1, cols=status_table.shape[1])
table.style = 'Table Grid'

# Add header row
for j in range(status_table.shape[1]):
    table.cell(0, j).text = status_table.columns[j]

# Add table data
for i in range(status_table.shape[0]):
    for j in range(status_table.shape[1]):
        table.cell(i + 1, j).text = str(status_table.iat[i, j])

# Create a plot for Submission Status and insert directly into the Word document
plt.figure(figsize=(4.5, 4.5))
sns.barplot(data=status_table_long, x='Year', y='Count', hue='Status')
plt.title('Submission Status Counts per Year')
plt.xlabel('Year')
plt.ylabel('Count')
plt.legend(title='Status', loc='upper right')

# Save plot to BytesIO
img_stream = BytesIO()
plt.savefig(img_stream, format='png')
plt.close()
img_stream.seek(0)

# Insert the plot directly into the Word document
doc.add_paragraph(
    "Graph: Submission Status Counts per Year. This bar chart shows the number of submissions classified by their status for each year from 2019 to 2023. "
    "Each bar represents a year, and the different colors correspond to different submission statuses. The graph allows for a comparison of submission status "
    "trends across the years, helping to identify whether certain submission statuses (such as 'Pending', 'Approved', etc.) have become more frequent or less so. "
    "This can provide valuable insights into areas that may require attention, such as delays in processing or changes in submission trends."
)
doc.add_picture(img_stream, width=Inches(4.5))

doc.add_paragraph('\n' * 2)

# Add Assessment Status Table to Word
doc.add_heading('Assessment Status Table by Year', level=1)

# Add paragraph explaining the table and graph
doc.add_paragraph(
    "The table below displays the assessment status counts for each year from 2019 to 2023, showing how many assessments were completed, pending, or in progress. "
    "The following graph visualizes the assessment status counts per year, helping to highlight trends and changes in assessment processing over time. "
    "By examining the graph, we can assess whether the number of completed assessments has increased or decreased, indicating improvements or challenges in the assessment process."
)

assessmentstatus_table = pd.DataFrame(assessmentstatus_counts)
assessmentstatus_table = assessmentstatus_table.T.fillna(0).astype(int)

# Add Assessment Status table using Table Grid style
table = doc.add_table(rows=assessmentstatus_table.shape[0] + 1, cols=assessmentstatus_table.shape[1])
table.style = 'Table Grid'

# Add header row
for j in range(assessmentstatus_table.shape[1]):
    table.cell(0, j).text = assessmentstatus_table.columns[j]

# Add table data
for i in range(assessmentstatus_table.shape[0]):
    for j in range(assessmentstatus_table.shape[1]):
        table.cell(i + 1, j).text = str(assessmentstatus_table.iat[i, j])

# Create a plot for Assessment Status and insert directly into the Word document
plt.figure(figsize=(4.5, 4.5))
sns.barplot(data=assessmentstatus_table_long_filtered, x='Year', y='Count', hue='Status')
plt.title('Assessment Status Counts per Year')
plt.xlabel('Year')
plt.ylabel('Count')
plt.legend(title='Status', loc='upper right')

# Save plot to BytesIO
img_stream = BytesIO()
plt.savefig(img_stream, format='png')
plt.close()
img_stream.seek(0)

# Insert the plot directly into the Word document
doc.add_paragraph(
    "Graph: Assessment Status Counts per Year. This bar chart illustrates the number of assessments completed, pending, or in progress for each year from 2019 to 2023. "
    "By observing the chart, we can identify whether assessment processing has become more efficient or if there have been delays or bottlenecks in certain years. "
    "Tracking assessment status over time is crucial for understanding the overall efficiency and timeliness of pension submissions."
)
doc.add_picture(img_stream, width=Inches(4.5))

doc.add_paragraph('\n' * 2)

# Add Approval Status Table to Word
doc.add_heading('Approval Status Table by Year', level=1)

# Add paragraph explaining the table and graph
doc.add_paragraph(
    "The table below shows the approval status counts for each year from 2019 to 2023. It highlights how many submissions were approved, rejected, or pending. "
    "The graph below visualizes these approval statuses over time, enabling a clear comparison of how approval rates have changed. "
    "This visualization helps to pinpoint trends in approval behavior, highlighting potential shifts in approval or rejection patterns."
)

approvalstatus_table = pd.DataFrame(approvalstatus_counts)
approvalstatus_table = approvalstatus_table.T.fillna(0).astype(int)

# Add Approval Status table using Table Grid style
table = doc.add_table(rows=approvalstatus_table.shape[0] + 1, cols=approvalstatus_table.shape[1])
table.style = 'Table Grid'

# Add header row
for j in range(approvalstatus_table.shape[1]):
    table.cell(0, j).text = approvalstatus_table.columns[j]

# Add table data
for i in range(approvalstatus_table.shape[0]):
    for j in range(approvalstatus_table.shape[1]):
        table.cell(i + 1, j).text = str(approvalstatus_table.iat[i, j])

# Create a plot for Approval Status and insert directly into the Word document
plt.figure(figsize=(4.5, 4.4))
sns.barplot(data=approvalstatus_table_long_filtered, x='Year', y='Count', hue='Status')
plt.title('Approval Status Counts per Year')
plt.xlabel('Year')
plt.ylabel('Count')
plt.legend(title='Status', loc='upper right')

# Save plot to BytesIO
img_stream = BytesIO()
plt.savefig(img_stream, format='png')
plt.close()
img_stream.seek(0)

# Insert the plot directly into the Word document
doc.add_paragraph(
    "Graph: Approval Status Counts per Year. This bar chart shows the number of approvals, rejections, and pending statuses for each year from 2019 to 2023. "
    "By analyzing the graph, we can identify any significant trends in approval or rejection rates, which can indicate the effectiveness of the approval process. "
    "A sharp increase or decrease in approvals could suggest changes in policy or submission quality, while a rise in rejections could signal issues that need addressing."
)
doc.add_picture(img_stream, width=Inches(4.5))

# Save the document

#import io
#doc_io = io.BytesIO()
#doc.save(doc_io)
#doc_io.seek(0)
#st.download_button(
 #   label="📄 Download Word Report",
  #  data=doc_io,
   # file_name="Summary_Statistics_2019_2023_with_detailed_graphs.docx",
    #mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
#)


# Filter the data based on 'approval_sts' equal to 'No' and 'year' equal to 2023
filtered_data2023 = pendata_combined[(pendata_combined['approval_sts'] == 'No') & (pendata_combined['year'] == 2023)]

# Select the specified columns for the new DataFrame
selected_columns = [
    'name_of_fund', 'approval_sts', 'missing_resolution', 
    'asset_split_diff', 'prior_year_non_approved', 'bonus_issues', 
    'asset_quality', 'other'
]

# Create the new DataFrame with the selected columns
new_dataframe2023 = filtered_data2023[selected_columns]


# Filter the data based on 'approval_sts' equal to 'No' and 'year' equal to 2023
filtered_data = pendata_combined[(pendata_combined['approval_sts'] == 'No') & (pendata_combined['year'] == 2019)]

# Select the specified columns for the new DataFrame
selected_columns = [
    'name_of_fund', 'approval_sts', 'missing_resolution', 
    'asset_split_diff', 'prior_year_non_approved', 'bonus_issues', 
    'asset_quality', 'other'
]

# Create the new DataFrame with the selected columns
new_dataframe19 = filtered_data[selected_columns]

# Filter the data based on 'approval_sts' equal to 'No' and 'year' equal to 2020
filtered_data = pendata_combined[(pendata_combined['approval_sts'] == 'No') & (pendata_combined['year'] == 2020)]

# Select the specified columns for the new DataFrame
selected_columns = [
    'name_of_fund', 'approval_sts', 'missing_resolution', 
    'asset_split_diff', 'prior_year_non_approved', 'bonus_issues', 
    'asset_quality', 'other'
]

# Create the new DataFrame with the selected columns
new_dataframe20 = filtered_data[selected_columns]

# Filter the data based on 'approval_sts' equal to 'No' and 'year' equal to 2021
filtered_data2021 = pendata_combined[(pendata_combined['approval_sts'] == 'pending') & (pendata_combined['year'] == 2021)]

# Select the specified columns for the new DataFrame
selected_columns = [
    'name_of_fund', 'approval_sts', 'missing_resolution', 
    'asset_split_diff', 'prior_year_non_approved', 'bonus_issues', 
    'asset_quality', 'other'
]

# Create the new DataFrame with the selected columns
new_dataframe2021 = filtered_data2021[selected_columns]


# Filter the data based on 'approval_sts' equal to 'No' and 'year' equal to 2022
filtered_data2022 = pendata_combined[(pendata_combined['approval_sts'] == 'pending') & (pendata_combined['year'] == 2022)]

# Select the specified columns for the new DataFrame
selected_columns = [
    'name_of_fund', 'approval_sts', 'missing_resolution', 
    'asset_split_diff', 'prior_year_non_approved', 'bonus_issues', 
    'asset_quality', 'other'
]

# Create the new DataFrame with the selected columns
new_dataframe2022 = filtered_data2022[selected_columns]

from docx import Document
from docx.shared import Inches


# Add a title or section heading to the Word document
doc.add_heading('Data for Funds with No Approval (2019-2023)', level=1)

# Add paragraph explaining the tables
doc.add_paragraph(
    "The tables below show the details of the funds for the years 2019 to 2023 with approval status marked as 'No'. "
    "They include information on various missing or pending factors that require attention such as missing resolutions, asset split discrepancies, "
    "prior year non-approvals, bonus issues, asset quality, and other issues."
)

# Loop through years 2019 to 2023 and filter the data
for year in range(2019, 2024):
    # Filter the data for each year with 'approval_sts' equal to 'No'
    filtered_data = pendata_combined[((pendata_combined['approval_sts'] == 'pending') | (pendata_combined['approval_sts'] == 'No')) & 
    (pendata_combined['year'] == year)
]


    # Select the specified columns for the new DataFrame
    selected_columns = [
        'name_of_fund', 'approval_sts', 'missing_resolution', 
        'asset_split_diff', 'prior_year_non_approved', 'bonus_issues', 
        'asset_quality', 'other'
    ]
    
    # Create the new DataFrame with the selected columns
    new_dataframe = filtered_data[selected_columns]

    # Add a heading for the specific year
    doc.add_heading(f'Data for Funds with Non Approval in {year}', level=2)

    # Add a paragraph explaining the table for the specific year
    doc.add_paragraph(
        f"The table below shows the details of the funds for the year {year} with approval status marked as 'No'. "
        "It includes information on various missing or pending factors such as missing resolutions, asset split discrepancies, "
        "prior year non-approvals, bonus issues, asset quality, and other issues."
    )

    # Add the DataFrame as a table to the Word document
    table = doc.add_table(rows=new_dataframe.shape[0] + 1, cols=new_dataframe.shape[1])
    table.style = 'Table Grid'

    # Add the header row (column names)
    for j, column_name in enumerate(new_dataframe.columns):
        table.cell(0, j).text = column_name

    # Add the table data (rows)
    for i in range(new_dataframe.shape[0]):
        for j in range(new_dataframe.shape[1]):
            table.cell(i + 1, j).text = str(new_dataframe.iat[i, j])

    # Add a line break between the years' data
    doc.add_paragraph('\n')

# Save the document

import io
doc_io = io.BytesIO()
doc.save(doc_io)
doc_io.seek(0)
st.download_button(
    label="📄 Download Word Report",
    data=doc_io,
    file_name="Summary_Statistics_2019_2023_with_detailed_graphs.docx",
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)

end_time = time.time()
run_time = end_time - start_time
print(f"Total run time: {run_time:.2f} seconds")

####
# Simulating a long-running task
with st.spinner('Processing data...'):
    time.sleep(5)  # Replace this with your actual processing logic
